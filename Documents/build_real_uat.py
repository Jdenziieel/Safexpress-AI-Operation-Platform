# -*- coding: utf-8 -*-
"""Build the REAL UAT for SafexpressOps.

Replaces the technical QA-style UAT-Questionnaire-v2.docx with a proper
User Acceptance Test:
  - Tester profile + baseline of current manual workflow
  - Scenario-based testing (8 Admin + 7 Manager + 6 User + 1 cross-role)
  - System Usability Scale (SUS) - industry standard
  - Business outcome survey (Likert + reasoning)
  - Formal acceptance / sign-off page

Output: Documents/UAT-Real.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"Documents\UAT-Real.docx"

CHECKBOX = "\u2610"  # empty ballot box

doc = Document()

# ----- page setup -----
section = doc.sections[0]
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ----- styles tuning -----
styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10.5)


# ============================================================
# helpers
# ============================================================
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    return p


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(text, italic=False, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    return p


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def add_kv_table(rows, header_color="DCE6F1", widths=(4.5, 12.5)):
    """rows = list of (label, value)."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        kc = t.rows[i].cells[0]
        vc = t.rows[i].cells[1]
        shade_cell(kc, header_color)
        kc.paragraphs[0].add_run(k).bold = True
        vc.paragraphs[0].add_run(str(v))
    set_col_widths(t, widths)
    return t


def add_blank():
    doc.add_paragraph()


def add_scenario(num, role, title, business_goal, objective, prerequisites, steps,
                 manual_baseline_label="Time it takes you today (manual): _____ min"):
    """Render a single test scenario."""
    add_h(f"{num}. {title}", level=3)

    # Quick metadata table
    add_kv_table([
        ("Role", role),
        ("Business goal", business_goal),
        ("Capstone objective", objective),
        ("Prerequisites", prerequisites),
    ])
    add_blank()

    # Steps (numbered)
    add_para("Steps the tester will perform:", bold=True)
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(s)
    add_blank()

    # Outcome / measurement table
    add_para("Tester records:", bold=True)
    out = doc.add_table(rows=6, cols=2)
    out.style = "Light Grid Accent 1"
    out.rows[0].cells[0].text = "Outcome"
    out.rows[0].cells[1].text = (
        f"{CHECKBOX} Pass (worked as expected)   "
        f"{CHECKBOX} Partial (worked but with issues)   "
        f"{CHECKBOX} Fail (did not work)"
    )
    out.rows[1].cells[0].text = "Manual baseline"
    out.rows[1].cells[1].text = manual_baseline_label
    out.rows[2].cells[0].text = "Time using the system"
    out.rows[2].cells[1].text = "Actual: _____ min"
    out.rows[3].cells[0].text = "Satisfaction (1-5)"
    out.rows[3].cells[1].text = (
        f"{CHECKBOX} 1 - Poor    "
        f"{CHECKBOX} 2 - Below avg    "
        f"{CHECKBOX} 3 - OK    "
        f"{CHECKBOX} 4 - Good    "
        f"{CHECKBOX} 5 - Excellent"
    )
    out.rows[4].cells[0].text = "Compared to your current process"
    out.rows[4].cells[1].text = (
        f"{CHECKBOX} Much better   {CHECKBOX} Better   {CHECKBOX} About the same   "
        f"{CHECKBOX} Worse   {CHECKBOX} Much worse"
    )
    out.rows[5].cells[0].text = "Comments / what went wrong / what would you change"
    out.rows[5].cells[1].text = "\n\n\n"
    for r in out.rows:
        shade_cell(r.cells[0], "DCE6F1")
        r.cells[0].paragraphs[0].runs[0].bold = True
    set_col_widths(out, (4.5, 12.5))
    add_blank()


# ============================================================
# COVER
# ============================================================
add_title("SafexpressOps")
add_subtitle("User Acceptance Test (Real UAT)")
add_blank()
add_para("Document type:  User Acceptance Test (business outcome validation)", bold=True)
add_para(
    "Companion document:  QA-FunctionalTestPlan.docx (technical functional test plan).",
    italic=True,
)
add_blank()
add_para(
    "This UAT measures whether SafexpressOps solves the operational problems described "
    "in the capstone's Chapter 1 - Project Context, not whether the code runs. Testers "
    "execute real Safexpress Logistics Inc. (SLI) workflows end-to-end and report (a) "
    "whether the workflow completed, (b) how long it took compared to the current "
    "manual process, (c) how well-designed the experience was, and (d) whether they "
    "would adopt the system in their daily work."
)
add_blank()
add_para(
    'Per the capstone\'s reference framing, UAT must answer two questions: '
    '"Does the software enable users to do their jobs?" and "Is the software '
    'designed well enough that they can do them effectively?" '
    'Every section of this document contributes to one of those two answers.',
    italic=True,
)

doc.add_page_break()

# ============================================================
# HOW TO USE
# ============================================================
add_h("How to use this document", level=1)
add_para(
    "1. Print this document or open it on a tablet next to your workstation. "
    "Do NOT preview the answers - the value of UAT is in capturing your honest, "
    "first-time experience."
)
add_para(
    "2. Complete Part 1 (Tester Profile + Baseline) before touching the system. "
    "The baseline minutes you record here is the comparison point for every scenario."
)
add_para(
    "3. Complete Part 2 (Scenarios) for your assigned role only. The cross-role scenario "
    "in Part 3 requires three testers (1 Admin + 1 Manager + 1 User) working together "
    "and is scheduled separately."
)
add_para(
    "4. Immediately after finishing the scenarios, complete Parts 4 and 5 (SUS + business "
    "survey) without going back to the system. These measure your overall impression."
)
add_para(
    "5. Return the signed Part 6 (Sign-off) page to the project lead within 2 business days."
)
add_blank()

add_h("Acceptance criteria", level=2)
add_kv_table([
    ("Pass", "All Pass+Partial scenarios divided by total scenarios >= 85% AND zero Fail "
            "on Critical scenarios (marked C in the scenario list) AND SUS score >= 68 "
            "(the industry-average benchmark)."),
    ("Conditional accept", "Pass+Partial >= 70% AND no Critical Fail AND SUS >= 60. The "
                           "tester documents specific conditions that must be met before "
                           "the system goes live."),
    ("Reject", "Pass+Partial < 70% OR a Critical scenario failed OR SUS < 60. The system "
               "is returned to the developers with the failure list before re-test."),
], widths=(4.5, 12.5))

doc.add_page_break()

# ============================================================
# PART 1 - TESTER PROFILE
# ============================================================
add_h("Part 1 - Tester Profile and Baseline", level=1)
add_para(
    "This section is filled out BEFORE you start using the system. It anchors every "
    "scenario in your real day-to-day work so the comparisons in Part 2 are meaningful."
)
add_blank()

add_h("1.1  Tester information", level=2)
add_kv_table([
    ("Name", ""),
    ("Role at SLI", ""),
    ("Years at SLI", ""),
    ("Warehouse / department", ""),
    ("Date of test", ""),
    ("UAT role assigned", f"{CHECKBOX} Administrator   {CHECKBOX} Manager   {CHECKBOX} User"),
    ("Test environment", f"{CHECKBOX} Desktop   {CHECKBOX} Laptop   {CHECKBOX} Tablet   {CHECKBOX} Phone"),
    ("Browser", ""),
], widths=(5, 12))

add_h("1.2  Baseline - your daily workflow today", level=2)
add_para(
    "Answer based on how things work today, WITHOUT SafexpressOps. Be honest - we are "
    "comparing the new system against your current reality, not against a perfect process."
)
add_blank()

add_kv_table([
    ("Q1.  Walk us through the one operational task you do most often. What's the "
     "trigger, what tools/files/people do you touch, what's the output?",
     "\n\n\n\n"),
    ("Q2.  How long does that task take you on a typical day (minutes)?", ""),
    ("Q3.  What's the SINGLE most frustrating step in that task?", "\n\n"),
    ("Q4.  How often do you have to ask a colleague for company information "
     "(per week)?", ""),
    ("Q5.  What % of your week is spent on repetitive copy-paste / manual "
     "data-entry work?", ""),
    ("Q6.  What's something you wish you could ask the system but don't have a "
     "tool for today?", "\n\n"),
], widths=(8.5, 8.5))

doc.add_page_break()

# ============================================================
# PART 2 - SCENARIOS
# ============================================================
add_h("Part 2 - Real-World Test Scenarios", level=1)
add_para(
    "Each scenario is mapped to a Specific Objective from Chapter 1 of the capstone. "
    "Scenarios marked 'C' (Critical) MUST pass for the system to be acceptable in any "
    "form - they validate the system's central promises."
)
add_para(
    "Capstone Objectives reference: SO1 = Knowledge Base + AI chat, SO2 = Analytical "
    "reports (ABC, OPR, Workload), SO3 = AI Personal Assistant (Google Workspace), "
    "SO4 = Dynamic data mapping, SO5 = Modular agentic workflow orchestration.",
    italic=True,
)
add_blank()

# ----------------- ADMIN -----------------
add_h("2.1  Administrator scenarios", level=2)
add_para("Complete this section if your assigned UAT role is Administrator.", italic=True)
add_blank()

add_scenario(
    "A1", "Administrator",
    "Onboard a new SLI manager  [C - Critical]",
    "Add a new manager account so they can use the system on Day 1 without IT escalation.",
    "Cross-cutting (Account management, RBAC)",
    "A real or test Google Workspace email for the new manager.",
    [
        "Sign in as Admin and go to Manage Accounts.",
        "Click Onboard New Account; enter the manager's full name, email, and select role = Manager.",
        "Save and confirm a success message.",
        "Send the new manager their sign-in link (or have them try directly with their Google account).",
        "Have the manager sign in and confirm they land on their home page (AI Assistant).",
        "Confirm the manager's sidebar does NOT show Accounts, Document Extraction, or Admin Activity Logs.",
    ],
    "Time it takes you today (manual onboarding): _____ min",
)

add_scenario(
    "A2", "Administrator",
    "Move a real warehouse SOP into the Knowledge Base  [C - Critical]",
    "Replace 'ask Mr. Cruz where to find the inbound checking rate' with an instant cited answer.",
    "SO1 - Knowledge Base + AI chat",
    "A real SLI warehouse SOP or company policy PDF (5-30 pages).",
    [
        "Open Document Extraction.",
        "Upload the PDF and wait for the parse to finish.",
        "Scroll through the AI-generated chunks; edit any chunk that is wrong or split awkwardly.",
        "Push the document to the Knowledge Base.",
        "Open SFX Bot in a new tab.",
        "Ask a question whose answer you KNOW is on a specific page of the SOP "
        "(e.g. 'What is the inbound checking rate at the VFP warehouse?').",
        "Verify SFX Bot's answer is correct, complete, and includes a citation back to the SOP.",
        "Click the citation - verify it opens the right document and the right page.",
    ],
    "Time today to find this answer manually (search binders / ask staff): _____ min",
)

add_scenario(
    "A3", "Administrator",
    "Set a per-user token quota and confirm enforcement",
    "Cost control - prevent any single user from burning the company's monthly LLM budget.",
    "Cross-cutting (Efficiency, Cost control - Chapter 3 NFR)",
    "A test account whose quota you can safely lower without disrupting production.",
    [
        "Open Token Quota Admin.",
        "Lower the test user's monthly token limit to a small value (e.g. 5,000 tokens).",
        "Sign in as that test user.",
        "Use the AI Assistant repeatedly until you cross the limit.",
        "Confirm the user sees a clear 'quota exceeded' message (not a stack trace or blank screen).",
        "As Admin, raise the test user's limit back to a normal value.",
        "Confirm the test user can immediately use the AI Assistant again without needing to refresh.",
    ],
    "Manual baseline: cost-control today is _____ (e.g. 'no enforcement' / 'spreadsheet review monthly')",
)

add_scenario(
    "A4", "Administrator",
    "Audit a manager's last-week activity for compliance",
    "Reconstruct what a user did, when, and why - in under 5 minutes.",
    "Cross-cutting (Security, Logging - Chapter 3 NFR)",
    "A manager who has been active in the system for at least 7 days.",
    [
        "Open Admin Activity Logs.",
        "Filter by the manager's email and a 7-day window.",
        "Confirm you can see, at minimum: AI Assistant prompts, dangerous actions approved/rejected, "
        "Knowledge Base queries, and any files uploaded.",
        "Pick one specific event you know happened (e.g. 'they sent an email to the VFP team last Tuesday') "
        "and confirm it appears in the log.",
        "Export the filtered log to CSV (or screenshot it for the audit pack).",
    ],
    "Time today to reconstruct a week of activity (across email, sheets, etc.): _____ min",
)

add_scenario(
    "A5", "Administrator",
    "Approve / reject a dangerous AI Assistant action  [C - Critical]",
    "The system must NEVER act dangerously without human review (Draft-Verify-Execute, Ch3).",
    "Cross-cutting (Reliability, Human-in-the-loop - Chapter 3 NFR)",
    "A manager who is about to run a workflow that will trigger a dangerous action "
    "(e.g. send an email or delete a calendar event).",
    [
        "Have the manager ask the AI Assistant: 'Send a happy weekend email to my team.'",
        "Open the Pending Tasks queue as Admin.",
        "Locate the proposed action; verify the To, Subject, Body, and any attachments are clearly visible.",
        "Reject the action.",
        "Confirm the email was NOT sent (check the manager's Sent folder).",
        "Confirm the manager is informed in their chat thread that the action was rejected.",
        "Repeat with a second dangerous action; this time, Approve.",
        "Confirm the action was actually executed (email sent, event created, etc.).",
    ],
    "Today's process for reviewing risky staff actions: _____ (e.g. 'no review' / 'after-the-fact via email logs')",
)

add_scenario(
    "A6", "Administrator",
    "Receive a budget threshold alert email",
    "Get notified BEFORE the platform overshoots its monthly LLM budget.",
    "Cross-cutting (Cost control - Chapter 3 NFR Efficiency)",
    "Access to the configured admin alert inbox (default admin@safexpressops.com).",
    [
        "Open the Admin Budget panel.",
        "Set monthly budget to a small test value (e.g. $1.00) and threshold to 80%.",
        "Use the AI Assistant heavily (or have a manager run a few workflows) until cumulative spend "
        "approaches $0.80.",
        "Re-open the Budget panel as Admin.",
        "Within ~2 minutes, confirm an email titled 'Monthly budget alert' arrives in the configured inbox.",
        "Open the Budget panel - verify it shows which user(s) and which model(s) drove the spend.",
        "Raise the budget back to a production value.",
    ],
    "Today's process for catching budget overruns: _____",
)

add_scenario(
    "A7", "Administrator",
    "Reactivate a previously deactivated account",
    "Bring a returning employee back online without losing their history.",
    "Cross-cutting (Account management, Continuity)",
    "An account that was previously deactivated.",
    [
        "Open Manage Accounts; filter to inactive users.",
        "Find the deactivated user and click Reactivate.",
        "Confirm a clear success message.",
        "Have the user sign in - confirm they can immediately access their normal pages.",
        "Confirm their previous chat threads / KB queries / report history are still accessible.",
    ],
    "Manual baseline: today's reactivation process: _____ min",
)

add_scenario(
    "A8", "Administrator",
    "Negative test - role boundary enforcement  [C - Critical]",
    "Confirm RBAC is a real permission boundary, not just a hidden menu.",
    "Cross-cutting (Security - Chapter 3 NFR)",
    "A test User-role account.",
    [
        "Sign out of Admin; sign in as the test User account.",
        "In the address bar, navigate directly to the URL Admin uses for Manage Accounts.",
        "Confirm you are blocked - either redirected or shown a clear 'Access Denied' page.",
        "Repeat with the URL for Admin Activity Logs and Document Extraction.",
        "Confirm all three are blocked. None should show a blank page or a stack trace.",
    ],
    "N/A - this is a guardrail test, no manual baseline.",
)

doc.add_page_break()

# ----------------- MANAGER -----------------
add_h("2.2  Manager scenarios", level=2)
add_para("Complete this section if your assigned UAT role is Manager.", italic=True)
add_blank()

add_scenario(
    "M1", "Manager",
    "Delivery Order automation end-to-end  [C - Critical]",
    "Replace 15-30 minutes of manual copy-paste from email to spreadsheet with a single chat turn. "
    "This is the canonical SLI VFP-warehouse workflow described in Chapter 1.",
    "SO3 (AI Assistant) + SO5 (orchestration) + scope: outbound delivery-order automation",
    "A real delivery-order email in your Gmail with a PDF attachment, AND a Google Sheet "
    "you currently use as a delivery tracker.",
    [
        "Open the AI Assistant.",
        "In a single chat message, type: 'Read my latest delivery order email from VFP, "
        "extract the line items, and append them to my [your Sheet name] tracker.'",
        "Watch the assistant work - it will search emails, parse the PDF, then push to your Sheet.",
        "Approve any dangerous actions when prompted (e.g. writing to the sheet).",
        "Open the Google Sheet directly - verify the new rows are there, in the right order, with no truncation.",
        "Spot-check 3-5 line items against the original PDF for accuracy.",
        "Time the full flow vs your current manual process.",
    ],
    "Manual baseline (today's end-to-end time for one delivery order): _____ min",
)

add_scenario(
    "M2", "Manager",
    "Generate today's One-Page Report and share it",
    "Replace the manual OPR-build-from-WMS-export with one upload + one share.",
    "SO2 (Analytical reports)",
    "A real WMS daily-export Excel/CSV file from this morning.",
    [
        "Open Analysis Reports -> One-Page Report.",
        "Upload the WMS export.",
        "Verify the date is auto-detected correctly.",
        "Skim the AI column mapping - confirm the columns line up with what you would expect "
        "(do NOT verify internal tier-by-tier matching; the question is just 'are the columns right?').",
        "Approve and write to Google Sheets.",
        "Open the resulting OPR sheet - verify the totals match your expectations.",
        "Share the link with one teammate via the AI Assistant ('Email the OPR link to [name]').",
    ],
    "Manual baseline (today's end-to-end OPR creation + share): _____ min",
)

add_scenario(
    "M3", "Manager",
    "Multi-step compound action - schedule a meeting + attach a doc",
    "Replace 5+ minutes of switching between Calendar, Drive, and Email with a single chat turn.",
    "SO5 (Modular agentic orchestration) + SO3",
    "An OPR sheet (created in M2 or pre-existing) in your Drive.",
    [
        "Open the AI Assistant in the same thread or a new one.",
        "Type a single message: 'Schedule a 30-minute Ops Review with the VFP team next Tuesday "
        "at 10 AM, attach the latest OPR sheet to the invite, and add a 3-bullet agenda in the description.'",
        "Approve any dangerous actions (the calendar event creation is a moderate action; the email "
        "to attendees may be dangerous).",
        "Open Google Calendar directly - verify the event exists with correct date/time/attendees.",
        "Open the event - verify the OPR link is in the description, and the agenda is present.",
        "Confirm with one attendee that they received the invite.",
    ],
    "Manual baseline (today's app-switching to set up the same meeting): _____ min",
)

add_scenario(
    "M4", "Manager",
    "Dynamic Mapping - fill a partner-supplied template",
    "Replace 30-60 minutes of manually rearranging Excel columns to match a partner's template.",
    "SO4 (Dynamic data mapping)",
    "A target template Google Sheet from a partner / client + your source data file (Excel / CSV) for this week.",
    [
        "Open Dynamic Mapping.",
        "Upload the source file.",
        "Pick or paste the link to the target Google Sheet.",
        "Review the AI-suggested column mappings.",
        "Adjust any wrong mappings using the UI (drag-and-drop or dropdown).",
        "Confirm and write.",
        "Open the target Sheet - verify the data landed in the correct columns and there's no shift.",
    ],
    "Manual baseline (today's Excel column-by-column copy time): _____ min",
)

add_scenario(
    "M5", "Manager",
    "SFX Bot - find a policy answer in under 2 minutes",
    "Replace 'ask a senior colleague' with a self-service answer + citation.",
    "SO1 (Knowledge Base + AI chat)",
    "An SOP or company policy already in the KB (uploaded by Admin in scenario A2).",
    [
        "Open SFX Bot.",
        "Ask a real question you would actually need to answer "
        "(e.g. 'What is the procedure for handling a damaged inbound pallet at VFP?').",
        "Read the answer; verify it's correct and complete.",
        "Click the citation - verify it opens the right document at the right page.",
        "Ask a follow-up question that depends on the previous one. Verify the bot keeps context.",
    ],
    "Manual baseline (today's average time to find a policy answer): _____ min",
)

add_scenario(
    "M6", "Manager",
    "Approval intervention - reject and redirect",
    "When the AI proposes the wrong thing, you can stop and redirect without losing progress.",
    "Cross-cutting (Reliability + Human-in-the-loop, Ch3 NFR)",
    "Set up a workflow where the AI's first proposed action will have an obvious factual error "
    "(e.g. ask the assistant to 'send a project update with this week's delivery count' "
    "while the count it has access to is intentionally stale).",
    [
        "Trigger the workflow as described.",
        "When the dangerous action is shown for approval, find the wrong fact in the proposed body.",
        "Reject the action.",
        "In the same thread, type: 'The numbers were wrong. Re-pull this week's delivery count "
        "from my OPR sheet and try again.'",
        "Verify the assistant adapts and re-proposes with corrected data.",
        "Confirm there's no need to re-explain the rest of the context.",
    ],
    "Manual baseline (today's process when a colleague drafts something with the wrong data): _____ min",
)

add_scenario(
    "M7", "Manager",
    "Continuity across days - resume a paused workflow",
    "A real assistant remembers what you were doing yesterday.",
    "SO1 (contextual memory) + SO5 (workflow continuity)",
    "Time the test across two calendar days (or two long-separated sessions).",
    [
        "Day 1: open the AI Assistant and ask: 'Plan a delivery briefing for Friday - "
        "check my calendar, see who's available, and draft an agenda.'",
        "When the assistant proposes the calendar event for approval, do NOT approve - close the browser.",
        "Day 2 (next day or after a long break): open the AI Assistant and find yesterday's thread.",
        "Type: 'Looks good - go ahead and create the event.'",
        "Verify the assistant resumes the workflow, knows which date and which attendees, and creates "
        "the event without re-asking you for any of that context.",
    ],
    "Manual baseline (today's process when you have to pick up an in-progress task tomorrow): _____ min",
)

doc.add_page_break()

# ----------------- USER -----------------
add_h("2.3  User scenarios", level=2)
add_para("Complete this section if your assigned UAT role is User.", italic=True)
add_blank()

add_scenario(
    "U1", "User",
    "SFX Bot - get an answer to a real warehouse-floor question  [C - Critical]",
    "Get an answer to an actual question that comes up during your shift, in under 2 minutes, with a citation.",
    "SO1 (Knowledge Base + AI chat)",
    "An SOP / procedure document already in the KB.",
    [
        "Open SFX Bot.",
        "Ask a question you've actually had to ask a colleague this past month "
        "(e.g. 'What do I do if my barcode scanner won't read a pallet?').",
        "Read the answer.",
        "Verify it is correct, specific to SLI, and not generic.",
        "Click the citation - verify it opens the right page.",
        "Ask a sensible follow-up. Verify continuity.",
    ],
    "Manual baseline (today's average time to get this answer from a colleague): _____ min",
)

add_scenario(
    "U2", "User",
    "ABC Analysis on a real inventory file  [C - Critical]",
    "Classify this month's inventory into A/B/C tiers without needing an analyst.",
    "SO2 (Analytical reports)",
    "A real inventory transactions Excel file (1 month of activity is enough).",
    [
        "Open Analysis Reports -> ABC Analysis.",
        "Upload the Excel file.",
        "Verify the months auto-detected match what's actually in the file.",
        "Review thresholds (defaults: A=80%, B=15%, C=5%); change if your team uses a different convention.",
        "Run the analysis.",
        "Read the result - confirm the items in tier A are recognizably your high-movers.",
        "Export to Google Sheets.",
        "Open the exported sheet; spot-check 3-5 items against your own knowledge of the inventory.",
    ],
    "Manual baseline (today's time for a manual ABC computation in Excel): _____ min",
)

add_scenario(
    "U3", "User",
    "Workload Analysis - plan today's pallet schedule",
    "Plan today's worker assignments without trial-and-error in Excel.",
    "SO2 + SLI workforce planning",
    "Today's expected delivery item list (item description, pallets, items per pallet) and worker count.",
    [
        "Open Workload Analysis.",
        "Enter each item's pallet count and items-per-pallet. Add at least 3 items.",
        "Set today's worker count.",
        "Click Calculate Workload.",
        "Read the result - confirm the per-phase breakdown (Inbound, Put-away, Picking, Outbound) "
        "matches your gut estimate within ~10%.",
        "Export the result to PDF for the floor supervisor.",
    ],
    "Manual baseline (today's planning time for the same set of items): _____ min",
)

add_scenario(
    "U4", "User",
    "Generate yesterday's One-Page Report",
    "Produce yesterday's daily report so your supervisor can see it before standup.",
    "SO2 (Analytical reports)",
    "Yesterday's WMS daily-export Excel/CSV file.",
    [
        "Open Analysis Reports -> One-Page Report.",
        "Upload the file.",
        "Verify the auto-detected date is yesterday.",
        "Confirm the AI column mapping looks right.",
        "Approve and write to Google Sheets.",
        "Open the resulting sheet - verify totals match what you'd expect from yesterday's activity.",
    ],
    "Manual baseline (today's time for a manual OPR build): _____ min",
)

add_scenario(
    "U5", "User",
    "Multi-turn drill-down with SFX Bot",
    "Investigate a topic the way you would with a senior colleague - one broad question, then 3 follow-ups.",
    "SO1 (KB + contextual memory) + Ch3 NFR Usability ('without formal training')",
    "KB has at least 3 SOPs covering related topics.",
    [
        "Open SFX Bot.",
        "Start broad: 'Tell me about the inbound checking process at VFP.'",
        "Read the answer; then follow up: 'What if the items don't match the manifest?'",
        "Follow up again: 'Who authorizes the discrepancy report?'",
        "Final follow-up: 'Show me the form template for that report.'",
        "Verify each follow-up understood the prior context (didn't restart the conversation).",
    ],
    "Manual baseline (today's process - back-and-forth chat / phone with a senior): _____ min",
)

add_scenario(
    "U6", "User",
    "Negative test - role boundary",
    "Confirm you cannot accidentally see Admin or Manager-only things.",
    "Cross-cutting (Security - Chapter 3 NFR RBAC)",
    "Your User-role account.",
    [
        "Try to navigate (via the address bar) to Manage Accounts, Document Extraction, "
        "and Admin Activity Logs.",
        "Confirm each is blocked with a clear 'Access Denied' page (not a blank screen, "
        "not a stack trace).",
        "Confirm SFX Bot, Analysis Reports, Dynamic Mapping, and Profile DO open normally for you.",
    ],
    "N/A - guardrail test.",
)

doc.add_page_break()

# ============================================================
# PART 3 - CROSS-ROLE COMPOUND
# ============================================================
add_h("Part 3 - Cross-role end-to-end scenario  [C - Critical]", level=1)
add_para(
    "This scenario requires three testers (1 Administrator + 1 Manager + 1 User) executing "
    "one real Monday morning at the VFP warehouse end-to-end. It validates every Specific "
    "Objective at once: SO1, SO2, SO3, SO4, SO5.",
    italic=True,
)
add_blank()

add_h("X1.  Real Monday at VFP - email to staffing schedule", level=3)
add_kv_table([
    ("Role", "1 Administrator + 1 Manager + 1 User (collaborative)"),
    ("Business goal",
     "One real warehouse Monday, executed across all three roles, in under 90 minutes."),
    ("Capstone objective", "SO1 + SO2 + SO3 + SO4 + SO5 (all)"),
    ("Prerequisites",
     "A real overnight delivery-order email in the Manager's Gmail with PDF attachment; "
     "a Drive folder shared between Manager and User; KB seeded with a few real SOPs."),
])
add_blank()

add_para("Steps the three testers will perform together:", bold=True)

add_para("\u2022 ADMIN (T+0 min): confirms the Manager and User have correct token quotas "
         "and that the User has access to the relevant Drive folder.", size=10.5)
add_para("\u2022 MANAGER (T+5 min): runs the Delivery Order automation flow (= scenario M1) - "
         "AI Assistant reads the email, extracts the PDF, pushes to the master tracker.", size=10.5)
add_para("\u2022 MANAGER (T+25 min): generates today's One-Page Report (= scenario M2) and "
         "schedules the Ops Review meeting for next Tuesday (= scenario M3).", size=10.5)
add_para("\u2022 USER (T+45 min): opens Workload Analysis, plugs in today's items from the "
         "master tracker, calculates the staffing requirement (= scenario U3).", size=10.5)
add_para("\u2022 USER (T+55 min): hits an unusual item type - opens SFX Bot, looks up the "
         "procedure (= scenario U5).", size=10.5)
add_para("\u2022 ADMIN (T+70 min): reviews the day's pending dangerous-actions queue, approves "
         "legitimate ones, rejects any anomalies (= scenario A5).", size=10.5)
add_para("\u2022 ADMIN (T+85 min): audits the day's activity log to confirm everything is "
         "traceable (= scenario A4).", size=10.5)
add_blank()

add_para("Group scoring (one row per role; teams discuss after):", bold=True)
gx = doc.add_table(rows=5, cols=4)
gx.style = "Light Grid Accent 1"
hdr = gx.rows[0].cells
hdr[0].text = "Role"
hdr[1].text = "Total time taken"
hdr[2].text = "Outcome"
hdr[3].text = "Where did it slow down?"
for i, c in enumerate(hdr):
    shade_cell(c, "DCE6F1")
    c.paragraphs[0].runs[0].bold = True
gx.rows[1].cells[0].text = "Admin"
gx.rows[2].cells[0].text = "Manager"
gx.rows[3].cells[0].text = "User"
gx.rows[4].cells[0].text = "Group total"
for r in range(1, 5):
    gx.rows[r].cells[2].text = (
        f"{CHECKBOX} Pass   {CHECKBOX} Partial   {CHECKBOX} Fail"
    )
set_col_widths(gx, (3, 3, 5, 6))
add_blank()

add_para("Manual baseline (today's time for the same Monday across the same three people): _____ min")
add_blank()

add_para("Reflection (group answer):", bold=True)
add_kv_table([
    ("Compared to today's process, the system would save us approximately:", "_____ min/day"),
    ("Single biggest friction point in the end-to-end flow:", "\n\n"),
    ("Single thing that delighted us:", "\n\n"),
    ("What would we have to fix BEFORE going live with this:", "\n\n"),
], widths=(7, 10))

doc.add_page_break()

# ============================================================
# PART 4 - SUS
# ============================================================
add_h("Part 4 - System Usability Scale (SUS)", level=1)
add_para(
    "Industry-standard 10-question survey. Score each statement on a 1-5 Likert scale "
    "(1 = Strongly Disagree, 5 = Strongly Agree). Do NOT spend more than 5 minutes on this - "
    "your immediate gut reaction is the right answer.",
    italic=True,
)
add_blank()

sus_statements = [
    "1.  I think that I would like to use SafexpressOps frequently.",
    "2.  I found SafexpressOps unnecessarily complex.",
    "3.  I thought SafexpressOps was easy to use.",
    "4.  I think that I would need the support of a technical person to use SafexpressOps.",
    "5.  I found the various functions in SafexpressOps were well integrated.",
    "6.  I thought there was too much inconsistency in SafexpressOps.",
    "7.  I would imagine that most people would learn to use SafexpressOps very quickly.",
    "8.  I found SafexpressOps very cumbersome to use.",
    "9.  I felt very confident using SafexpressOps.",
    "10. I needed to learn a lot of things before I could get going with SafexpressOps.",
]

sus_t = doc.add_table(rows=11, cols=6)
sus_t.style = "Light Grid Accent 1"
hdr = sus_t.rows[0].cells
hdr[0].text = "Statement"
hdr[1].text = "1 - SD"
hdr[2].text = "2 - D"
hdr[3].text = "3 - N"
hdr[4].text = "4 - A"
hdr[5].text = "5 - SA"
for c in hdr:
    shade_cell(c, "DCE6F1")
    c.paragraphs[0].runs[0].bold = True
for i, st in enumerate(sus_statements, 1):
    row = sus_t.rows[i].cells
    row[0].text = st
    for j in range(1, 6):
        row[j].text = CHECKBOX
        row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
set_col_widths(sus_t, (8.5, 1.7, 1.7, 1.7, 1.7, 1.7))
add_blank()

add_para(
    "SUS scoring (the project lead computes this; the tester does not have to):",
    italic=True, size=9,
)
add_para(
    "  - For odd-numbered items: subtract 1 from the score.\n"
    "  - For even-numbered items: subtract the score from 5.\n"
    "  - Multiply the sum by 2.5 to get a 0-100 score.\n"
    "  - SUS >= 68 is industry-average; >= 80 is excellent; < 50 is unusable.",
    italic=True, size=9,
)

doc.add_page_break()

# ============================================================
# PART 5 - BUSINESS OUTCOME SURVEY
# ============================================================
add_h("Part 5 - Business outcome survey", level=1)
add_para(
    "These six questions answer whether the system, beyond being usable, is actually "
    "valuable to your work and to SLI as a business.",
    italic=True,
)
add_blank()

bus_statements = [
    "B1.  I believe SafexpressOps would save me significant time on my recurring tasks.",
    "B2.  I would trust the system's outputs (reports, mapped data, drafted emails) for my real work.",
    "B3.  I would adopt SafexpressOps in my daily work if it were available tomorrow.",
    "B4.  I could learn to use SafexpressOps effectively WITHOUT formal training.",
    "B5.  I would recommend SafexpressOps to my SLI colleagues.",
    "B6.  Compared to my current process, SafexpressOps is a meaningful improvement (not just a different way to do the same thing).",
]

bus_t = doc.add_table(rows=7, cols=6)
bus_t.style = "Light Grid Accent 1"
hdr = bus_t.rows[0].cells
hdr[0].text = "Statement"
hdr[1].text = "1 - SD"
hdr[2].text = "2 - D"
hdr[3].text = "3 - N"
hdr[4].text = "4 - A"
hdr[5].text = "5 - SA"
for c in hdr:
    shade_cell(c, "DCE6F1")
    c.paragraphs[0].runs[0].bold = True
for i, st in enumerate(bus_statements, 1):
    row = bus_t.rows[i].cells
    row[0].text = st
    for j in range(1, 6):
        row[j].text = CHECKBOX
        row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
set_col_widths(bus_t, (8.5, 1.7, 1.7, 1.7, 1.7, 1.7))
add_blank()

add_h("Open feedback", level=2)
add_kv_table([
    ("Top 3 things you LIKED about SafexpressOps:", "1.\n\n2.\n\n3.\n\n"),
    ("Top 3 things that FRUSTRATED you while testing:", "1.\n\n2.\n\n3.\n\n"),
    ("ONE feature that, if added or improved, would make this indispensable for your job:", "\n\n"),
    ("Anything you would NOT trust the system to do (and why)?", "\n\n"),
], widths=(7.5, 9.5))

doc.add_page_break()

# ============================================================
# PART 6 - SIGN-OFF
# ============================================================
add_h("Part 6 - Formal Acceptance / Sign-off", level=1)
add_para(
    "This page produces the documented go/no-go milestone for SafexpressOps. By signing "
    "below, you state that you have personally completed the scenarios assigned to your "
    "role and that your verdict reflects your honest first-time experience.",
    italic=True,
)
add_blank()

add_h("6.1  Acceptance verdict", level=2)
verdict_t = doc.add_table(rows=3, cols=2)
verdict_t.style = "Light Grid Accent 1"
verdict_t.rows[0].cells[0].text = f"{CHECKBOX} ACCEPT"
verdict_t.rows[0].cells[1].text = (
    "All Critical [C] scenarios passed. Pass+Partial >= 85%. SUS >= 68. "
    "I accept the system as ready for production deployment."
)
verdict_t.rows[1].cells[0].text = f"{CHECKBOX} ACCEPT WITH CONDITIONS"
verdict_t.rows[1].cells[1].text = (
    "All Critical [C] scenarios passed. Pass+Partial >= 70%. SUS >= 60. I accept "
    "the system on the condition that the issues listed in 6.2 are addressed before go-live."
)
verdict_t.rows[2].cells[0].text = f"{CHECKBOX} REJECT"
verdict_t.rows[2].cells[1].text = (
    "One or more Critical [C] scenarios failed, OR Pass+Partial < 70%, OR SUS < 60. "
    "The system is not ready for production. The failures listed in 6.2 must be fixed "
    "and the system re-tested."
)
for r in verdict_t.rows:
    shade_cell(r.cells[0], "DCE6F1")
    r.cells[0].paragraphs[0].runs[0].bold = True
set_col_widths(verdict_t, (5, 12))
add_blank()

add_h("6.2  Conditions / Required fixes (if Accept-with-Conditions or Reject)", level=2)
add_kv_table([
    ("Condition #1", "\n\n"),
    ("Condition #2", "\n\n"),
    ("Condition #3", "\n\n"),
    ("Condition #4", "\n\n"),
    ("Condition #5", "\n\n"),
], widths=(4, 13))

add_h("6.3  Tester signature", level=2)
add_kv_table([
    ("Name (printed)", ""),
    ("Role at SLI", ""),
    ("UAT role assigned", ""),
    ("Signature", ""),
    ("Date", ""),
], widths=(4, 13))

add_h("6.4  Witness / Project lead signature", level=2)
add_kv_table([
    ("Name (printed)", ""),
    ("Role", ""),
    ("Signature", ""),
    ("Date", ""),
], widths=(4, 13))

doc.add_page_break()

# ============================================================
# APPENDIX
# ============================================================
add_h("Appendix - How this UAT maps to the capstone", level=1)
add_para(
    "Traceability matrix - every scenario in this document maps back to a Specific Objective "
    "from Chapter 1 of the capstone. This table is what Chapter 4's UAT Analysis section "
    "should cite when reporting acceptance results.",
    italic=True,
)
add_blank()

trace_rows = [
    ("Capstone objective", "Tested by scenarios"),
    ("SO1 - Knowledge Base + AI chat (SFX Bot)", "A2, M5, U1, U5, X1"),
    ("SO2 - Analytical reports (ABC, OPR, Workload)", "M2, U2, U3, U4, X1"),
    ("SO3 - AI Personal Assistant (Google Workspace)", "M1, M3, M6, X1"),
    ("SO4 - Dynamic data mapping engine", "M4, X1"),
    ("SO5 - Modular agentic workflow orchestration", "M3, M7, X1"),
    ("Cross-cutting - Account management & RBAC", "A1, A7, A8, U6"),
    ("Cross-cutting - Cost control / Token quotas", "A3, A6"),
    ("Cross-cutting - Reliability / Human-in-the-loop", "A5, M6"),
    ("Cross-cutting - Logging & Audit", "A4, X1"),
]
tt = doc.add_table(rows=len(trace_rows), cols=2)
tt.style = "Light Grid Accent 1"
for i, (k, v) in enumerate(trace_rows):
    tt.rows[i].cells[0].text = k
    tt.rows[i].cells[1].text = v
    if i == 0:
        for c in tt.rows[i].cells:
            shade_cell(c, "DCE6F1")
            c.paragraphs[0].runs[0].bold = True
set_col_widths(tt, (8.5, 8.5))
add_blank()

add_para(
    "End of UAT-Real.docx. Companion document: QA-FunctionalTestPlan.docx (for Chapter 4's "
    "Functional Requirements section).",
    italic=True, size=9,
)

doc.save(OUT)
print(f"OK -> {OUT}")
