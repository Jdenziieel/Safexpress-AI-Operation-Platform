# -*- coding: utf-8 -*-
"""Build the LITE (scale-based) UAT for SafexpressOps.

A standalone questionnaire-only UAT with no walkthrough scenarios. The
tester gets a 20-30 minute hands-on session with the system, then fills
out this form (15 minutes). Designed for sample sizes of 15-25 SLI staff.

Sections:
  Part 1 - Tester profile + current-workflow baseline
  Part 2 - System Usability Scale (SUS) - 5-point Likert
  Part 3 - Business outcome scale - 5-point Likert
  Part 4 - Open-ended feedback
  Part 5 - Formal sign-off
  Appendix - Traceability to capstone Specific Objectives

Terminology note: this document uses "scale" throughout (Likert scale,
satisfaction scale, agreement scale). The word "rating" is avoided.

Output: Documents/UAT-Lite.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"Documents\UAT-Lite.docx"
CHECKBOX = "\u2610"

doc = Document()

section = doc.sections[0]
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_h(text, level=1):
    doc.add_heading(text, level=level)


def add_para(text, italic=False, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if size:
        r.font.size = Pt(size)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, c in enumerate(row.cells):
            if i < len(widths_cm):
                c.width = Cm(widths_cm[i])


def add_kv_table(rows, widths=(4.5, 12.5), header_color="DCE6F1"):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        kc = t.rows[i].cells[0]
        vc = t.rows[i].cells[1]
        shade_cell(kc, header_color)
        kc.paragraphs[0].add_run(k).bold = True
        vc.paragraphs[0].add_run(str(v))
    set_col_widths(t, widths)
    return t


def add_blank():
    doc.add_paragraph()


def add_likert_table(statements, widths=(8.5, 1.7, 1.7, 1.7, 1.7, 1.7)):
    """Render a 5-point Likert scale table (1=SD, 2=D, 3=N, 4=A, 5=SA)."""
    t = doc.add_table(rows=len(statements) + 1, cols=6)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Statement"
    hdr[1].text = "1 - SD"
    hdr[2].text = "2 - D"
    hdr[3].text = "3 - N"
    hdr[4].text = "4 - A"
    hdr[5].text = "5 - SA"
    for c in hdr:
        shade_cell(c, "DCE6F1")
        c.paragraphs[0].runs[0].bold = True
    for i, s in enumerate(statements, 1):
        row = t.rows[i].cells
        row[0].text = s
        for j in range(1, 6):
            row[j].text = CHECKBOX
            row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_col_widths(t, widths)
    return t


# ============================================================
# COVER
# ============================================================
add_title("SafexpressOps")
add_subtitle("User Acceptance Test - Scale-Based Questionnaire")
add_blank()
add_para("Document type:  User Acceptance Test (Likert scale survey)", bold=True)
add_para(
    "Companion document:  QA-FunctionalTestPlan.docx (technical functional test plan).",
    italic=True,
)
add_blank()
add_para(
    "This UAT measures user acceptance of SafexpressOps using a Likert-scale "
    "questionnaire administered AFTER each tester has completed a hands-on session with "
    "the system. The instrument captures (a) general usability via the System Usability "
    "Scale (SUS), (b) business-outcome agreement via a 6-item scale tied to the capstone's "
    "Specific Objectives, and (c) open-ended feedback for thematic analysis."
)
add_blank()
add_para(
    'Per the capstone\'s reference framing, UAT must answer two questions: "Does the '
    'software enable users to do their jobs?" and "Is the software designed well enough '
    'that they can do them effectively?" The scales in Parts 2 and 3 are designed so that '
    'the aggregate scores directly answer those two questions, respectively.',
    italic=True,
)
add_blank()

add_h("Administration protocol", level=2)
add_para(
    "1. Each tester receives a 20-30 minute hands-on session with SafexpressOps before "
    "completing this form. The session must let the tester actually use the system - "
    "a screen-share demo alone does not constitute valid UAT exposure."
)
add_para(
    "2. Testers complete this form INDEPENDENTLY (no group filling, no live coaching). "
    "Each form takes about 15 minutes to complete."
)
add_para(
    "3. The project lead collects all forms, computes aggregate scores, and reports the "
    "results in Chapter 4 - Test Case and User Acceptance Test Analysis."
)
add_blank()

add_h("Acceptance criteria", level=2)
add_kv_table([
    ("Sample size",
     "Minimum 10 testers, target 15-25, drawn from the three SLI roles "
     "(Administrator, Manager, User) in approximately the same proportion they "
     "appear in the workforce."),
    ("Pass",
     "Mean SUS score >= 68 (industry-average benchmark) AND mean business-outcome "
     "scale score >= 3.5/5 across all 6 items AND ZERO items with mean < 3.0."),
    ("Conditional pass",
     "Mean SUS score 60-67 OR business-outcome mean 3.0-3.49. The system is accepted "
     "with the requirement that the lowest-scoring items are addressed before go-live."),
    ("Reject",
     "Mean SUS < 60 OR business-outcome mean < 3.0 OR any single item with mean < 2.5. "
     "The system is returned to the developers."),
], widths=(4.5, 12.5))

doc.add_page_break()

# ============================================================
# PART 1 - TESTER PROFILE & BASELINE
# ============================================================
add_h("Part 1 - Tester profile and baseline", level=1)
add_para(
    "Filled BEFORE you start using the system. The information here lets us slice the "
    "results by role and warehouse so that any patterns in the scale scores can be "
    "interpreted in context."
)
add_blank()

add_h("1.1  Tester information", level=2)
add_kv_table([
    ("Name (or anonymized ID)", ""),
    ("Role at SLI", ""),
    ("Years at SLI", ""),
    ("Warehouse / department", ""),
    ("Date of test", ""),
    ("Hands-on session duration", "_____ minutes"),
    ("Test environment",
     f"{CHECKBOX} Desktop   {CHECKBOX} Laptop   {CHECKBOX} Tablet   {CHECKBOX} Phone"),
    ("Browser", ""),
], widths=(5, 12))

add_h("1.2  Workflow baseline (current process, no SafexpressOps)", level=2)
add_para(
    "Six short questions about your daily work TODAY. Answer based on how things "
    "actually run, not how they should run. This anchors the scale scores in Part 3.",
    italic=True,
)
add_blank()

add_kv_table([
    ("Q1.  Which one operational task do you do most often? Briefly describe the "
     "trigger, the tools/files/people you touch, and the output.",
     "\n\n\n"),
    ("Q2.  How long does that task usually take you (minutes)?", ""),
    ("Q3.  What is the SINGLE most frustrating step in that task?", "\n\n"),
    ("Q4.  How often per week do you have to ask a colleague for company "
     "information you cannot find yourself?", ""),
    ("Q5.  What % of your week is spent on repetitive copy-paste / manual "
     "data-entry work?", ""),
    ("Q6.  What is something you wish you could ask the system, but no tool "
     "does today?", "\n\n"),
], widths=(8.5, 8.5))

doc.add_page_break()

# ============================================================
# PART 2 - SUS
# ============================================================
add_h("Part 2 - System Usability Scale (SUS)", level=1)
add_para(
    "Industry-standard 10-item scale used in academic UAT research for over 30 years. "
    "Mark one box per row using the 5-point Likert scale (1 = Strongly Disagree, "
    "5 = Strongly Agree). Trust your first reaction - do not go back to the system "
    "to re-check.",
    italic=True,
)
add_blank()

sus = [
    "1.  I think that I would like to use SafexpressOps frequently.",
    "2.  I found SafexpressOps unnecessarily complex.",
    "3.  I thought SafexpressOps was easy to use.",
    "4.  I think that I would need the support of a technical person to use SafexpressOps.",
    "5.  I found the various functions in SafexpressOps were well integrated.",
    "6.  I thought there was too much inconsistency in SafexpressOps.",
    "7.  I would imagine that most people would learn to use SafexpressOps very quickly.",
    "8.  I found SafexpressOps very cumbersome to use.",
    "9.  I felt very confident using SafexpressOps.",
    "10. I needed to learn a lot of things before I could get going with SafexpressOps.",
]
add_likert_table(sus)
add_blank()

add_para(
    "SUS scoring (computed by the project lead, not the tester):",
    italic=True, size=9,
)
add_para(
    "  - Odd-numbered items: subtract 1 from the score.\n"
    "  - Even-numbered items: subtract the score from 5.\n"
    "  - Sum the converted scores and multiply by 2.5 to get a 0-100 score.\n"
    "  - SUS >= 68 is industry-average; >= 80 is excellent; < 50 is unusable.",
    italic=True, size=9,
)

doc.add_page_break()

# ============================================================
# PART 3 - BUSINESS OUTCOME SCALE
# ============================================================
add_h("Part 3 - Business outcome scale", level=1)
add_para(
    "Six items mapped to the capstone's Specific Objectives (see Appendix A). Use the "
    "same 5-point Likert scale (1 = Strongly Disagree, 5 = Strongly Agree). Optional: "
    "in the brief reasoning rows below the table, write one sentence explaining why "
    "you scored an item at the extreme (1, 2, or 5). Reasoning is not required.",
    italic=True,
)
add_blank()

bus = [
    "B1.  I believe SafexpressOps would save me significant time on my recurring tasks.",
    "B2.  I would trust the system's outputs (reports, mapped data, drafted emails) for my real work.",
    "B3.  I would adopt SafexpressOps in my daily work if it were available tomorrow.",
    "B4.  I could learn to use SafexpressOps effectively WITHOUT formal training.",
    "B5.  I would recommend SafexpressOps to my SLI colleagues.",
    "B6.  Compared to my current process, SafexpressOps is a meaningful improvement (not just a different way to do the same thing).",
]
add_likert_table(bus)
add_blank()

add_h("3.1  Optional reasoning for any extreme scores (1, 2, or 5)", level=2)
add_kv_table([
    ("Item B1 reasoning (if scored 1, 2, or 5)", "\n"),
    ("Item B2 reasoning (if scored 1, 2, or 5)", "\n"),
    ("Item B3 reasoning (if scored 1, 2, or 5)", "\n"),
    ("Item B4 reasoning (if scored 1, 2, or 5)", "\n"),
    ("Item B5 reasoning (if scored 1, 2, or 5)", "\n"),
    ("Item B6 reasoning (if scored 1, 2, or 5)", "\n"),
], widths=(7, 10))

doc.add_page_break()

# ============================================================
# PART 4 - OPEN-ENDED FEEDBACK
# ============================================================
add_h("Part 4 - Open-ended feedback", level=1)
add_para(
    "Free-text answers, not scale-based. The project lead will perform thematic "
    "analysis on these responses to surface recurring issues that the Likert scales "
    "alone may miss.",
    italic=True,
)
add_blank()

add_kv_table([
    ("Q1.  Top 3 things you LIKED about SafexpressOps:", "1.\n\n2.\n\n3.\n\n"),
    ("Q2.  Top 3 things that FRUSTRATED you while using SafexpressOps:",
     "1.\n\n2.\n\n3.\n\n"),
    ("Q3.  ONE feature that, if added or improved, would make this system "
     "indispensable for your job:", "\n\n"),
    ("Q4.  Anything you would NOT trust the system to do automatically? Why?",
     "\n\n"),
    ("Q5.  Any other feedback for the project team?", "\n\n"),
], widths=(7.5, 9.5))

doc.add_page_break()

# ============================================================
# PART 5 - FORMAL SIGN-OFF
# ============================================================
add_h("Part 5 - Formal acceptance / sign-off", level=1)
add_para(
    "By signing below, you confirm that you completed a hands-on session with "
    "SafexpressOps and that the scale scores and feedback above reflect your honest "
    "first-time experience.",
    italic=True,
)
add_blank()

add_h("5.1  Personal verdict", level=2)
verdict_t = doc.add_table(rows=3, cols=2)
verdict_t.style = "Light Grid Accent 1"
verdict_t.rows[0].cells[0].text = f"{CHECKBOX} ACCEPT"
verdict_t.rows[0].cells[1].text = (
    "I would use SafexpressOps in my daily work and recommend it to my colleagues."
)
verdict_t.rows[1].cells[0].text = f"{CHECKBOX} ACCEPT WITH CONDITIONS"
verdict_t.rows[1].cells[1].text = (
    "I would use SafexpressOps once the issues I noted above are addressed."
)
verdict_t.rows[2].cells[0].text = f"{CHECKBOX} REJECT"
verdict_t.rows[2].cells[1].text = (
    "I would NOT use SafexpressOps in its current form. The reasons are documented in Part 4."
)
for r in verdict_t.rows:
    shade_cell(r.cells[0], "DCE6F1")
    r.cells[0].paragraphs[0].runs[0].bold = True
set_col_widths(verdict_t, (5, 12))
add_blank()

add_h("5.2  Signature", level=2)
add_kv_table([
    ("Name (printed)", ""),
    ("Role at SLI", ""),
    ("Signature", ""),
    ("Date", ""),
], widths=(4, 13))

add_h("5.3  Witness / project lead signature", level=2)
add_kv_table([
    ("Name (printed)", ""),
    ("Role", ""),
    ("Signature", ""),
    ("Date", ""),
], widths=(4, 13))

doc.add_page_break()

# ============================================================
# APPENDIX
# ============================================================
add_h("Appendix A - Mapping of scale items to capstone Specific Objectives", level=1)
add_para(
    "This is the traceability table used when reporting results in Chapter 4. Each "
    "Specific Objective from Chapter 1 is covered by at least one Likert item across "
    "Parts 2 and 3.",
    italic=True,
)
add_blank()

trace = [
    ("Capstone objective", "Covered by scale item(s)"),
    ("Cross-cutting (Usability, learnability)", "SUS items 1, 3, 4, 7, 9, 10"),
    ("Cross-cutting (Reliability, consistency)", "SUS items 2, 5, 6, 8"),
    ("SO1 - Knowledge Base + AI chat", "B2 (trust outputs), B4 (no training)"),
    ("SO2 - Analytical reports (ABC, OPR, Workload)", "B1 (time savings), B2 (trust)"),
    ("SO3 - AI Personal Assistant (Google Workspace)", "B1, B2, B3 (adopt tomorrow)"),
    ("SO4 - Dynamic data mapping engine", "B1, B2"),
    ("SO5 - Modular agentic workflow orchestration", "B6 (meaningful improvement)"),
    ("Adoption / change-management", "B3, B5 (recommend), Part 5 verdict"),
]
tt = doc.add_table(rows=len(trace), cols=2)
tt.style = "Light Grid Accent 1"
for i, (k, v) in enumerate(trace):
    tt.rows[i].cells[0].text = k
    tt.rows[i].cells[1].text = v
    if i == 0:
        for c in tt.rows[i].cells:
            shade_cell(c, "DCE6F1")
            c.paragraphs[0].runs[0].bold = True
set_col_widths(tt, (8.5, 8.5))
add_blank()

add_h("Appendix B - Reporting template for Chapter 4", level=1)
add_para(
    "When reporting results, use this structure (suggested):",
    italic=True,
)
add_blank()
add_para(
    "  - Demographics:  N = ___ testers (Admin: ___, Manager: ___, User: ___). Mean "
    "years at SLI = ___. Mean hands-on session duration = ___ min."
)
add_para(
    "  - SUS results:  Mean SUS = ___ / 100 (interpreted as: ___). Per-item means "
    "and standard deviations in Table ___."
)
add_para(
    "  - Business-outcome scale:  Mean across 6 items = ___ / 5. Highest-scoring item: "
    "B__ (mean ___). Lowest-scoring item: B__ (mean ___). Per-item table in Table ___."
)
add_para(
    "  - Open-ended thematic analysis:  Top 3 recurring positive themes were ___. "
    "Top 3 recurring frustration themes were ___."
)
add_para(
    "  - Verdict distribution:  Accept = ___% , Accept-with-Conditions = ___%, "
    "Reject = ___%."
)
add_para(
    "  - Acceptance decision:  Per the criteria in the cover page, the overall "
    "verdict is ___ (Pass / Conditional / Reject)."
)
add_blank()

add_para(
    "End of UAT-Lite.docx. Use UAT-Real.docx (the scenario-based version) only for "
    "deep-tester sessions; use this Lite version for the broad-sample survey.",
    italic=True, size=9,
)

doc.save(OUT)
print(f"OK -> {OUT}")
