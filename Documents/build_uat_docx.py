"""Render UAT-Questionnaire-v2.md to a Word .docx using python-docx.

Pragmatic markdown subset:
  '# X'   -> doc heading 1 (also pagebreaks at role headers)
  '## X'  -> heading 2
  '### X' -> heading 3
  '#### X'-> heading 4
  '> ...' -> indented italic note paragraph
  '|...|' -> table (header detected by following '|---|' row)
  '---'   -> blank divider
  '**x**' -> bold run
  '`x`'   -> Courier run
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path(r"c:\Users\Denz\Documents\tigers\Ai-Agents\Documents\UAT-Questionnaire-v2.md")
OUT = Path(r"c:\Users\Denz\Documents\tigers\Ai-Agents\Documents\UAT-Questionnaire-v2.docx")


# ---- helpers ----------------------------------------------------------

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_runs(p, text):
    """Add inline runs honoring **bold** and `code` tokens."""
    if not text:
        return
    text = text.replace("->", "\u2192")
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        else:
            r = p.add_run(part)


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="999999"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_pagebreak(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)


def configure_styles(doc):
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)

    for name, size, color in [
        ("Heading 1", 22, "1F4E78"),
        ("Heading 2", 15, "1F4E78"),
        ("Heading 3", 12, "2E75B6"),
        ("Heading 4", 11, "2E75B6"),
    ]:
        try:
            st = doc.styles[name]
            st.font.name = "Calibri"
            st.font.size = Pt(size)
            st.font.color.rgb = RGBColor.from_string(color)
            st.font.bold = True
        except KeyError:
            pass


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    add_runs(p, text)
    for r in p.runs:
        r.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def render_table(doc, header, body):
    n = len(header)
    is_qa = n == 4 and "Yes" in header and "No" in header
    table = doc.add_table(rows=len(body) + 1, cols=n)
    table.autofit = False
    table.allow_autofit = False

    if is_qa:
        widths = [Cm(1.0), Cm(13.5), Cm(1.5), Cm(1.5)]
    elif n == 2:
        widths = [Cm(4.5), Cm(13.0)]
    elif n == 3:
        widths = [Cm(4.5), Cm(7.0), Cm(6.0)]
    else:
        avail = 17.5
        widths = [Cm(avail / n)] * n

    for i, c in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        cell.width = widths[i]
        p = cell.paragraphs[0]
        if is_qa and i in (0, 2, 3):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(c)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        shade(cell, "1F4E78")
        set_cell_borders(cell, "1F4E78")

    for ri, row in enumerate(body, start=1):
        zebra = "F7F9FC" if ri % 2 == 0 else "FFFFFF"
        while len(row) < n:
            row.append("")
        for ci, val in enumerate(row[:n]):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            cell.width = widths[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            shade(cell, zebra)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if is_qa and ci in (0, 2, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if is_qa and ci in (2, 3) and not val.strip():
                add_runs(p, "\u2610")  # ballot box
            else:
                add_runs(p, val)
            for r in p.runs:
                r.font.size = Pt(10)


# ---- markdown parser --------------------------------------------------

def parse_table(lines, i):
    rows = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        rows.append(lines[i].strip())
        i += 1
    if len(rows) < 2:
        return None, i
    sep = rows[1]
    if not re.match(r"^\|\s*[-:|\s]+\|$", sep):
        return None, i

    def split_row(s):
        s = s.strip().strip("|")
        return [c.strip() for c in s.split("|")]

    header = split_row(rows[0])
    body = [split_row(r) for r in rows[2:]]
    return (header, body), i


def build():
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6)
        s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.6)
        s.right_margin = Inches(0.6)
    configure_styles(doc)

    role_seen = False
    title_done = False

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            content = stripped[2:].strip()
            if content in ("ADMIN", "MANAGER", "USER"):
                if role_seen:
                    add_pagebreak(doc)
                role_seen = True
                p = doc.add_heading(level=1)
                add_runs(p, content)
                i += 1
                continue
            if not title_done:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(content)
                r.bold = True
                r.font.size = Pt(20)
                r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
                title_done = True
            else:
                p = doc.add_heading(level=1)
                add_runs(p, content)
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_heading(level=2)
            add_runs(p, stripped[3:].strip())
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_heading(level=3)
            add_runs(p, stripped[4:].strip())
            i += 1
            continue

        if stripped.startswith("#### "):
            p = doc.add_heading(level=4)
            add_runs(p, stripped[5:].strip())
            i += 1
            continue

        if stripped.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                t = lines[i].lstrip()[1:].lstrip()
                buf.append(t)
                i += 1
            text_block = " ".join(b for b in buf if b)
            if text_block:
                add_quote(doc, text_block)
            continue

        if stripped.startswith("|"):
            parsed, ni = parse_table(lines, i)
            if parsed:
                header, body = parsed
                render_table(doc, header, body)
                doc.add_paragraph()
                i = ni
                continue

        para = []
        while i < len(lines) and lines[i].strip() and not lines[i].lstrip().startswith(("#", ">", "|")):
            if lines[i].strip() == "---":
                break
            para.append(lines[i].strip())
            i += 1
        if para:
            p = doc.add_paragraph()
            add_runs(p, " ".join(para))

    doc.save(str(OUT))
    print("WROTE:", OUT, OUT.stat().st_size, "bytes")


if __name__ == "__main__":
    build()
