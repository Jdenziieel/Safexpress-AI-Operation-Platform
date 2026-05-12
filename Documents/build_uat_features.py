# -*- coding: utf-8 -*-
"""Build the feature-based UAT for SafexpressOps (v3 - Yes/No functional).

CHANGES vs v2:
  - Removed Part 1 (Tester profile + baseline).
  - Functional acceptance (was Part 3, now Part 2) converted from 5-point
    Likert to Yes/No. SUS, Quality, and Adoption sections stay Likert.
  - Open-ended feedback (was Part 6, now Part 5) trimmed from 6 to 3
    questions.
  - Added 3 items to plug gaps identified during the Chapter 1 scope audit:
      F3.8 - AI Assistant real-time progress visibility (WebSocket scope item)
      F3.x - Email label management folded into F3.1
      F8.2 - PDF preview / page nav / side-by-side highlight overlay
        (Document Extraction UI scope item)
  - New Appendix A: Coverage audit. Cross-references every Chapter 1 scope
    item to the UAT item that tests it (or marks it 'back-end, not
    user-testable' with rationale).
  - Renumbered all parts after Part 1 was removed.

Output: Documents/UAT-FeatureScale.docx (overwrites the previous build)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os
PRIMARY = r"Documents\UAT-FeatureScale.docx"
FALLBACK = r"Documents\UAT-FeatureScale-v3.docx"
try:
    with open(PRIMARY, "ab") as _t:
        pass
    OUT = PRIMARY
except PermissionError:
    OUT = FALLBACK
CHECKBOX = "\u2610"

doc = Document()
sect = doc.sections[0]
sect.left_margin = Cm(1.8)
sect.right_margin = Cm(1.8)
sect.top_margin = Cm(2.0)
sect.bottom_margin = Cm(2.0)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_h(text, level=1):
    doc.add_heading(text, level=level)


def add_para(text, italic=False, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if size:
        r.font.size = Pt(size)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, c in enumerate(row.cells):
            if i < len(widths_cm):
                c.width = Cm(widths_cm[i])


def add_kv_table(rows, widths=(4.5, 13.0)):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        kc = t.rows[i].cells[0]
        vc = t.rows[i].cells[1]
        shade_cell(kc, "DCE6F1")
        kc.paragraphs[0].add_run(k).bold = True
        vc.paragraphs[0].add_run(str(v))
    set_col_widths(t, widths)
    return t


def add_blank():
    doc.add_paragraph()


def add_yesno_table(items, widths=(1.0, 13.0, 1.7, 1.7)):
    """4-column Yes/No table: # | Question | Yes | No.

    Mirrors the format used in the team's earlier UAT-Questionnaire-v2.docx
    so testers familiar with that doc see a consistent interface.
    """
    t = doc.add_table(rows=len(items) + 1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Question"
    hdr[2].text = "Yes"
    hdr[3].text = "No"
    for c in hdr:
        shade_cell(c, "DCE6F1")
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, q in enumerate(items, 1):
        row = t.rows[i].cells
        row[0].text = str(i)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = q
        row[2].text = CHECKBOX
        row[3].text = CHECKBOX
        for j in (2, 3):
            row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_col_widths(t, widths)
    return t


def add_likert_table(items, widths=(8.5, 1.4, 1.4, 1.4, 1.4, 1.4, 1.5)):
    t = doc.add_table(rows=len(items) + 1, cols=7)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Statement"
    hdr[1].text = "1 - SD"
    hdr[2].text = "2 - D"
    hdr[3].text = "3 - N"
    hdr[4].text = "4 - A"
    hdr[5].text = "5 - SA"
    hdr[6].text = "N/A"
    for c in hdr:
        shade_cell(c, "DCE6F1")
        c.paragraphs[0].runs[0].bold = True
    for i, s in enumerate(items, 1):
        row = t.rows[i].cells
        row[0].text = s
        for j in range(1, 7):
            row[j].text = CHECKBOX
            row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_col_widths(t, widths)
    return t


def role_tag(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
    r.font.size = Pt(9.5)


# ============================================================
# COVER
# ============================================================
add_title("SafexpressOps")
add_subtitle("User Acceptance Test - Feature-Based Questionnaire")
add_blank()
add_para("Document type:  User Acceptance Test (mixed Yes/No + Likert)", bold=True)
add_para(
    "Companion documents:  QA-FunctionalTestPlan.docx (technical functional test plan), "
    "UAT-Real.docx (scenario-based deep UAT).",
    italic=True,
)
add_blank()

add_para(
    "This UAT measures user acceptance of every shipped feature in SafexpressOps. "
    "Functional acceptance items are scored Yes/No (binary outcome - did the "
    "feature work as expected?). Usability, quality, and adoption items are "
    "scored on a 5-point Likert scale (degree of agreement). Each tester "
    "completes the form after a 30-45 minute hands-on session with the system."
)
add_blank()

add_para(
    'Per the capstone\'s reference framing, UAT must answer two questions: "Does '
    'the software enable users to do their jobs?" and "Is the software designed '
    'well enough that they can do them effectively?" Part 2 (Functional Yes/No) '
    'directly answers the first. Parts 1 (SUS), 3 (Quality), and 4 (Adoption) '
    'together answer the second.',
    italic=True,
)
add_blank()

add_h("Why this is feature-based, not requirement-based", level=2)
add_para(
    "The capstone's Chapter 3 Functional / Non-Functional Requirements list "
    "was the original specification, not the final shipped feature set. Some "
    "requirements were superseded, some were merged, and a few new features "
    "were added during development. This UAT is anchored on the actual "
    "shipped system (features visible to a tester in the running app) so "
    "that test results are decisive: a high score means the shipped product "
    "works for users, not that the original specification was satisfied. "
    "Traceability back to Chapter 1 Specific Objectives AND scope items is "
    "preserved in Appendix A and Appendix B."
)
add_blank()

add_h("Why functional items are Yes/No, not Likert", level=2)
add_para(
    "Functional items measure binary outcomes - the export either contained "
    "all four tabs or it did not; the calendar event either has the right "
    "attendees or it does not. Forcing a 5-point scale on a binary outcome "
    "(\"4-Agree that the event was correct\") is methodologically weak and "
    "dilutes the signal. The Yes/No format also matches the team's earlier "
    "UAT-Questionnaire-v2.docx so testers see a familiar interface. The "
    "remaining sections (SUS, Quality, Adoption) measure SUBJECTIVE "
    "experience - those use Likert because \"I felt confident\" or "
    "\"I would recommend\" are not binary."
)
add_blank()

add_h("Administration protocol", level=2)
add_para(
    "1. Tester receives a 30-45 minute hands-on session covering the "
    "features in their role's scope (Admin: all sections; Manager: all "
    "except 2.8 and 2.9; User: all except 2.3, 2.8, and 2.9). A demo-only "
    "walkthrough does NOT count - the tester must actually use the system."
)
add_para(
    "2. Tester completes this form independently. Time required: ~25-35 min."
)
add_para(
    "3. Project lead aggregates all forms, computes per-section Yes-rate "
    "and Likert means + standard deviations, and reports per Appendix C."
)
add_blank()

add_h("Acceptance criteria", level=2)
add_kv_table([
    ("Sample size", "Minimum 10 testers (target 15-25), distributed across "
                    "the three SLI roles (Administrator, Manager, User) in "
                    "approximately the same ratio as the actual workforce."),
    ("Pass",
     "Every functional sub-section's Yes-rate >= 85% AND SUS mean >= 68 "
     "AND Adoption scale mean >= 3.5/5 AND zero individual functional items "
     "with Yes-rate < 70%."),
    ("Conditional pass",
     "Functional sub-section Yes-rates 70-84% OR SUS mean 60-67 OR "
     "Adoption mean 3.0-3.49. Accepted with the requirement that the "
     "lowest-scoring items are fixed before go-live."),
    ("Reject",
     "Any functional sub-section Yes-rate < 70% OR SUS mean < 60 OR "
     "Adoption mean < 3.0 OR any individual functional item with "
     "Yes-rate < 50%."),
])

doc.add_page_break()

# ============================================================
# PART 1 - SUS (was Part 2)
# ============================================================
add_h("Part 1 - System Usability Scale (SUS)", level=1)
add_para(
    "Industry-standard 10-item Likert scale used in academic UAT for over "
    "30 years. Trust your first reaction. The N/A column should NOT be "
    "used here - every SUS item applies to every tester.",
    italic=True,
)
add_blank()

sus_items = [
    "1.  I think that I would like to use SafexpressOps frequently.",
    "2.  I found SafexpressOps unnecessarily complex.",
    "3.  I thought SafexpressOps was easy to use.",
    "4.  I think that I would need the support of a technical person to use SafexpressOps.",
    "5.  I found the various functions in SafexpressOps were well integrated.",
    "6.  I thought there was too much inconsistency in SafexpressOps.",
    "7.  I would imagine that most people would learn to use SafexpressOps very quickly.",
    "8.  I found SafexpressOps very cumbersome to use.",
    "9.  I felt very confident using SafexpressOps.",
    "10. I needed to learn a lot of things before I could get going with SafexpressOps.",
]
add_likert_table(sus_items)
add_blank()
add_para(
    "SUS scoring (computed by the project lead): odd items - subtract 1; "
    "even items - subtract from 5; sum and multiply by 2.5 to get a 0-100 "
    "score. >= 68 industry average; >= 80 excellent; < 50 unusable.",
    italic=True, size=9,
)

doc.add_page_break()

# ============================================================
# PART 2 - FUNCTIONAL (Yes/No) - was Part 3
# ============================================================
add_h("Part 2 - Functional acceptance (Yes / No)", level=1)
add_para(
    "Each feature area uses a Yes/No format: tick Yes if the feature "
    "behaved as the question describes during your hands-on session. If "
    "the feature does not apply to your role (see role-applicability tag "
    "under each heading), leave both boxes blank. The project lead will "
    "compute the Yes-rate per item and per section.",
    italic=True,
)
add_blank()

# 2.1 Authentication
add_h("2.1  Authentication, roles, and access", level=2)
role_tag("Applies to: ALL roles (Administrator, Manager, User)")
add_yesno_table([
    "Were you able to sign in to SafexpressOps using your Google account without needing technical help?",
    "Did the role-based access correctly match your job - you were not exposed to features outside your role, and you had access to everything you needed?",
    "Could you find and read your profile, role badge, and current token / quota usage easily?",
    "Did logging out and logging back in work cleanly, with your chat history and saved work preserved?",
])
add_blank()

# 2.2 SFX Bot
add_h("2.2  SFX Bot - knowledge base question-answering", level=2)
role_tag("Applies to: ALL roles")
add_yesno_table([
    "Did SFX Bot answer your company-knowledge questions accurately when the answer existed in the knowledge base?",
    "Did SFX Bot's responses include citations (filename, page) that pointed to the correct source document?",
    "When you asked something NOT in the knowledge base, did SFX Bot say so instead of guessing?",
    "Did SFX Bot keep the context of your conversation across multiple back-and-forth messages in the same thread?",
    "Could you create, switch between, rename, and delete chat sessions, with only your own sessions visible to you?",
    "Did SFX Bot stream responses in real time, and could you stop a streaming response if you no longer needed it?",
])
add_blank()

# 2.3 AI Assistant
add_h("2.3  AI Assistant - personal productivity (Google Workspace)", level=2)
role_tag("Applies to: Administrator + Manager only.  Users: skip this section.")
add_yesno_table([
    "Did the AI Assistant correctly perform the Gmail tasks you asked of it (draft, send, reply, forward, search, manage labels)?",
    "Did the AI Assistant correctly perform the Google Calendar tasks you asked of it (create, view, update, delete events; invite attendees; generate Meet links)?",
    "Did the AI Assistant correctly perform the Google Drive tasks you asked of it (upload files, create / name folders, list, search, download)?",
    "Did the AI Assistant correctly perform the Google Docs / Sheets tasks you asked of it (create, edit, read, format)?",
    "Did the AI Assistant successfully complete multi-step workflows (e.g. 'find the latest delivery-order email, extract the PDF, push to my sheet') in a single chat turn?",
    "Were dangerous actions (sending emails, deleting events, etc.) always shown to you for explicit approval before they ran?",
    "After you approved or rejected an action, did the AI Assistant continue the workflow correctly without losing earlier context?",
    "Within the same chat session, did the AI Assistant retain the context of your conversation - so you did not have to re-explain earlier instructions or repeat information already provided in the same thread?",
    "Could you see real-time progress updates while the AI Assistant worked - including the current stage (Analyzing / Planning / Executing / Composing), the step / tool name, the step counter (e.g. Step 2 of 3), and the elapsed time?",
])
add_blank()

# 2.4 Dynamic Mapping
add_h("2.4  Dynamic Mapping", level=2)
role_tag("Applies to: ALL roles")
add_yesno_table([
    "Could you upload Excel (.xlsx) and CSV source files, with the system rejecting unsupported types (.txt, .png, etc.) and corrupt files using clear error messages?",
    "When your source Excel had multiple sheets, did the system auto-detect them, auto-select the best match when one sheet clearly dominated, and show a sheet-picker when the choice was ambiguous?",
    "When your source contained multiple stacked data sections (title row + header row + data rows, repeated), did the system detect them and let you select the correct section?",
    "Could you paste a Google Sheets URL into the target field, see all tabs listed (with the URL's gid pre-selected if present), and have tab names with spaces still resolved correctly?",
    "Did the AI propose accurate column mappings - matching exact column names without an AI call, and using AI for synonyms (e.g. 'Qty' to 'Quantity') - and correctly choose the row-anchor strategy (date, entity, composite key, label/value, horizontal, or cross_tab) for the target layout?",
    "Before writing, did the system show a preview with the cell-level diff (current value vs new value), updated-row vs appended-row counts, and formula columns excluded from the diff - and could you deselect individual rows to skip them?",
    "After you confirmed, did the system write into the correct rows / columns / sections, report the number of rows updated and appended, preserve formula cells, and refuse to write when all rows had been deselected?",
    "Could you save and reload named column-mapping templates, with only your own (or your team's) templates visible?",
    "When something went wrong (no columns matched, Google auth error, zero rows updated), did the system surface a clear actionable message instead of a stack trace?",
])
add_blank()

# 2.5 ABC Analysis
add_h("2.5  ABC Analysis", level=2)
role_tag("Applies to: ALL roles")
add_yesno_table([
    "Could you upload an Excel (.xlsx) transaction file with sensible defaults pre-filled (Date column = Transdate, Item column = Itemcode, Quantity = Qtyordered, plus Description and UOM), and have non-Excel / empty files rejected with a clear message?",
    "Could you change the Date / Item / Quantity / Description / UOM column names to match your file's headers - and if you entered a column that did not exist, did the system clearly report the missing column instead of crashing?",
    "Could you set custom Class A and Class B cumulative thresholds (default 70% / 90%) and have the system warn you or block running when A-threshold was greater than or equal to B-threshold?",
    "Did the system automatically detect all months present in the date column without manual configuration, label them correctly (e.g. 'Jan 2025'), and run successfully even when the file covered only a single month?",
    "Did the classification result give each item exactly one class (A / B / C), rank items by combined score (quantity multiplied by order count), show monotonically increasing cumulative percentages, and surface all required per-item columns (Rank, Item Code, Description, UOM, Total Qty, Order Count, Item Score, Percentage, Cumulative Pct, Class) plus a monthly comparison and an Executive Summary?",
    "After running, was a new Google Sheet auto-created with all required tabs (Executive Summary, Monthly Comparison, Complete ABC Analysis, Class A / B / C, plus one tab per detected month), all openable without permission errors via the provided link?",
    "Could you view a history list of your past ABC analyses, re-open any past result, and only see your own (not other users') analyses?",
])
add_blank()

# 2.6 OPR
add_h("2.6  One-Page Report (OPR)", level=2)
role_tag("Applies to: ALL roles")
add_yesno_table([
    "Could you upload an Excel (.xlsx) or CSV daily-operations file, with PDFs / images / corrupt files rejected via a specific error message and the filename + size shown so you could confirm the correct file?",
    "Did the system correctly find and parse the Date column, recognise common date formats (Excel-style, ISO, US, dotted), skip blank or unparseable rows gracefully, and report the count of dates extracted (e.g. '31 dates extracted')?",
    "Did the SmartMappingEngine produce accurate mappings - via Tier 0 forced typo corrections (e.g. 'Toal Manhours' to 'Total Manhours'), Tier 1 exact match, Tier 2 Levenshtein similarity, Tier 3 OpenAI fallback - while preserving formula columns, never confusing Inbound vs Outbound or Manhours vs Safe-Manhours, and showing a confidence score per source-target column pair?",
    "Did the mapping preview list every column with its proposed target and confidence, flag low-confidence mappings for review, allow you to override or 'Skip' any mapping via a dropdown, show a cell-level diff (current vs new value) excluding formula cells, and report matched-vs-unmatched date counts?",
    "After you confirmed, did the system write the data into the 'DATA ENTRY' sheet using exactly your approved mappings (no AI re-run), update only matching dates with cell-level writes that preserve formula cells in the same row, and report the number of rows and cells updated?",
    "When errors occurred (no source columns mapped to the target, zero matching dates between source and sheet, Google auth error, header-only file with no data rows), did the system surface a clear actionable message - including a sample of the file's dates for the zero-date-match case?",
    "Could you view a history list of your past OPR processings (filename, date processed, rows updated), re-open any past record to review what was written, and only see your own (not other users') OPR records?",
])
add_blank()

# 2.7 Workload Analysis
add_h("2.7  Workload Analysis", level=2)
role_tag("Applies to: ALL roles")
add_yesno_table([
    "Did the Workload Analysis page load with at least one item row by default, and could you add new item rows, fill in description / pallets / items per pallet, and remove rows (with at least one row always remaining)?",
    "Could you expand the warehouse rate settings, see the default phase rates loaded, and edit the rates for Inbound Checking, Put-away, Picking, and Outbound Checking?",
    "When you clicked Calculate Workload with valid data, did the system produce totals (time, pallets, items, workers) plus a per-phase breakdown, with totals matching the inputs (total pallets equals entered pallet count; total items equals pallets multiplied by items per pallet)?",
    "If required data was missing (no item data or empty workers field), did the system stop the calculation and show a clear message?",
    "Could you export the result to PDF with the summary, phase breakdown, and items breakdown all included?",
    "Could you save the calculation to the database when the backend was connected, with the Save button correctly disabled when no result existed or when the backend was offline - and could you still calculate using the loaded default rates while offline, with the offline state communicated clearly?",
])
add_blank()

# 2.8 KB Management
add_h("2.8  Knowledge Base management", level=2)
role_tag("Applies to: Administrator only.  Manager + User: skip this section.")
add_yesno_table([
    "Could you upload a PDF, see it parsed into chunks with proper structure (headings, tables, embedded images), and edit chunks if needed before pushing to the knowledge base?",
    "Did the Document Extraction interface show a real-time PDF preview with page navigation and side-by-side highlight overlay marking where each chunk came from on the page?",
    "Did the system detect duplicate uploads (by content hash and filename) and warn you before re-processing?",
    "Could you see KB analytics (queries per document, token usage, popular topics) in a clear dashboard?",
    "After you pushed a new document, was SFX Bot able to answer questions about it and cite it correctly?",
])
add_blank()

# 2.9 Account & Activity
add_h("2.9  Account and activity management", level=2)
role_tag("Applies to: Administrator only.  Manager + User: skip this section.")
add_yesno_table([
    "Could you onboard new accounts (set name, email, role) and have the new user sign in immediately?",
    "Could you edit, deactivate, and reactivate accounts cleanly, with chat history preserved across reactivation?",
    "Could you review activity logs by user, time window, or action type to support audit / compliance?",
    "Could you set per-user token quotas and confirm they were enforced when the user crossed the limit?",
    "Could you review and approve or reject pending dangerous actions from a single queue?",
])

doc.add_page_break()

# ============================================================
# PART 3 - QUALITY (Likert)
# ============================================================
add_h("Part 3 - Quality scale (Reliability, Security, Compatibility, Efficiency)", level=1)
add_para(
    "These items measure how well the system behaves overall, in the "
    "tester's perception. Maps to the Chapter 3 Non-Functional "
    "Requirements that are testable from a user's perspective. The "
    "remaining NFRs (Maintainability, Scalability) are assessed via "
    "developer-side code review and load tests respectively, not through "
    "this UAT.",
    italic=True,
)
add_blank()

quality_items = [
    "Q1.  The system gave me clear, actionable error messages when something went wrong - not technical jargon or stack traces.",
    "Q2.  When the system was unsure about my request, it asked me to clarify rather than guessing.",
    "Q3.  Dangerous actions never ran without my explicit approval.",
    "Q4.  I felt confident that my data and credentials were handled securely.",
    "Q5.  I was only able to see and act on what my role allows. I was never exposed to other users' data or features outside my role.",
    "Q6.  SafexpressOps worked properly on the device and browser I normally use at work.",
    "Q7.  The system responded fast enough that I did not lose my train of thought waiting for it.",
    "Q8.  Quota / usage indicators were clear, so I never ran out of capacity unexpectedly.",
]
add_likert_table(quality_items)

doc.add_page_break()

# ============================================================
# PART 4 - ADOPTION (Likert)
# ============================================================
add_h("Part 4 - Adoption and business outcome scale", level=1)
add_para(
    "These items measure whether you would actually USE the system. The "
    "aggregate score here is the most important signal of whether "
    "SafexpressOps is ready to deploy at SLI.",
    italic=True,
)
add_blank()

adoption_items = [
    "B1.  I believe SafexpressOps would save me significant time on my recurring tasks.",
    "B2.  I would trust the system's outputs (reports, mapped data, drafted emails) for my real work.",
    "B3.  I would adopt SafexpressOps in my daily work if it were available tomorrow.",
    "B4.  I could learn to use SafexpressOps effectively WITHOUT formal training.",
    "B5.  I would recommend SafexpressOps to my SLI colleagues.",
    "B6.  Compared to my current process, SafexpressOps is a meaningful improvement (not just a different way to do the same thing).",
]
add_likert_table(adoption_items)

doc.add_page_break()

# ============================================================
# PART 5 - OPEN-ENDED (trimmed to 3)
# ============================================================
add_h("Part 5 - Open-ended feedback", level=1)
add_para(
    "Free-text answers. The project lead will run thematic analysis "
    "across all testers' responses to surface recurring issues that the "
    "Yes/No and Likert items alone cannot capture.",
    italic=True,
)
add_blank()
add_kv_table([
    ("Q1.  Top 3 things you LIKED about SafexpressOps:", "1.\n\n2.\n\n3.\n\n"),
    ("Q2.  Top 3 things that FRUSTRATED you while using SafexpressOps:",
     "1.\n\n2.\n\n3.\n\n"),
    ("Q3.  ONE feature that, if added or improved, would make SafexpressOps "
     "indispensable for your job:", "\n\n"),
])

doc.add_page_break()

# ============================================================
# PART 6 - SIGN-OFF
# ============================================================
add_h("Part 6 - Formal acceptance / sign-off", level=1)
add_para(
    "By signing below, you confirm that you completed a hands-on session "
    "with SafexpressOps (not a demo only) and that the answers above "
    "reflect your honest first-time experience.",
    italic=True,
)
add_blank()

add_h("6.1  Personal verdict", level=2)
verdict_t = doc.add_table(rows=3, cols=2)
verdict_t.style = "Light Grid Accent 1"
verdict_t.rows[0].cells[0].text = f"{CHECKBOX} ACCEPT"
verdict_t.rows[0].cells[1].text = (
    "I would use SafexpressOps in my daily work and recommend it to my colleagues."
)
verdict_t.rows[1].cells[0].text = f"{CHECKBOX} ACCEPT WITH CONDITIONS"
verdict_t.rows[1].cells[1].text = (
    "I would use SafexpressOps once the issues I noted above are addressed."
)
verdict_t.rows[2].cells[0].text = f"{CHECKBOX} REJECT"
verdict_t.rows[2].cells[1].text = (
    "I would NOT use SafexpressOps in its current form. The reasons are documented in Part 5."
)
for r in verdict_t.rows:
    shade_cell(r.cells[0], "DCE6F1")
    r.cells[0].paragraphs[0].runs[0].bold = True
set_col_widths(verdict_t, (5, 12.5))
add_blank()

add_h("6.2  Tester signature", level=2)
add_kv_table([
    ("Name (printed)", ""),
    ("Role at SLI", ""),
    ("UAT role assigned", ""),
    ("Signature", ""),
    ("Date", ""),
])

add_h("6.3  Witness / project lead signature", level=2)
add_kv_table([
    ("Name (printed)", ""),
    ("Role", ""),
    ("Signature", ""),
    ("Date", ""),
])

doc.add_page_break()

# ============================================================
# APPENDIX A - SPECIFIC OBJECTIVE TRACEABILITY
# ============================================================
add_h("Appendix A - Traceability to Chapter 1 Specific Objectives", level=1)
add_para(
    "Maps each section of this UAT to the Specific Objective from Chapter "
    "1 that it validates. Every SO has at least one section covering it.",
    italic=True,
)
add_blank()

so_rows = [
    ("UAT section", "Capstone Specific Objective"),
    ("Part 2.2 - SFX Bot",
     "SO1 - Knowledge Base + AI chat with contextual memory"),
    ("Part 2.8 - KB Management",
     "SO1 (admin side - upload, chunking, KB curation)"),
    ("Part 2.5 - ABC Analysis",
     "SO2 - Automated analytical reports"),
    ("Part 2.6 - OPR",
     "SO2 - Automated analytical reports"),
    ("Part 2.7 - Workload Analysis",
     "SO2 - Automated analytical reports"),
    ("Part 2.3 - AI Assistant",
     "SO3 - AI Personal Assistant (Google Workspace)"),
    ("Part 2.4 - Dynamic Mapping",
     "SO4 - Dynamic data mapping engine"),
    ("Part 2.3 (multi-step compound items F3.5 / F3.7 / F3.8)",
     "SO5 - Modular agentic workflow orchestration"),
    ("Part 2.1 - Auth, Roles, Access",
     "Cross-cutting (Security, RBAC)"),
    ("Part 2.9 - Account & Activity",
     "Cross-cutting (Security, Audit, Cost)"),
    ("Part 1 - SUS",
     "NFR Usability"),
    ("Part 3 - Quality",
     "NFR Reliability + Security + Compatibility + Efficiency"),
    ("Part 4 - Adoption",
     "Cross-cutting (overall acceptance signal)"),
]
tt = doc.add_table(rows=len(so_rows), cols=2)
tt.style = "Light Grid Accent 1"
for i, (a, b) in enumerate(so_rows):
    tt.rows[i].cells[0].text = a
    tt.rows[i].cells[1].text = b
    if i == 0:
        for cc in tt.rows[i].cells:
            shade_cell(cc, "DCE6F1")
            cc.paragraphs[0].runs[0].bold = True
set_col_widths(tt, (8.0, 9.5))

doc.add_page_break()

# ============================================================
# APPENDIX B - SCOPE COVERAGE AUDIT (NEW)
# ============================================================
add_h("Appendix B - Chapter 1 scope coverage audit", level=1)
add_para(
    "Every scope item from Chapter 1 of the capstone is listed below, with "
    "a pointer to the UAT item that tests it (or a rationale if the item "
    "is not user-testable). This is the table the defense panel can use to "
    "verify nothing in the scope was left untested.",
    italic=True,
)
add_blank()

scope_rows = [
    ("Chapter 1 scope item", "Tested by", "Notes"),

    # KB / PDF processing
    ("Upload PDF documents for processing and extraction", "Part 2.8 Q1", ""),
    ("Extract text with structure preservation (headings, paragraphs, lists)",
     "Part 2.8 Q1", "Tester confirms structure visible in chunk preview"),
    ("Detect and extract tables from PDFs", "Part 2.8 Q1",
     "Tester confirms tables appear in chunks"),
    ("Extract embedded images via PyMuPDF (Base64)", "Part 2.8 Q1",
     "Tester confirms images appear in chunks"),
    ("Intelligent chunking via GPT-4.1 / GPT-4o-mini", "Part 2.8 Q1",
     "User-observable as 'chunks make sense'"),
    ("Vector embeddings + Weaviate storage", "Part 2.8 Q5",
     "Indirectly tested - 'SFX Bot can answer about new doc'"),

    # SFX Bot / RAG
    ("Conversational Q&A (SFX Bot) with hybrid search (vector + keyword)",
     "Part 2.2 Q1", ""),
    ("Multi-turn conversations with context retention", "Part 2.2 Q4", ""),
    ("Markdown responses with inline citations", "Part 2.2 Q2",
     "Citation accuracy explicitly tested"),
    ("Token usage tracking per session and total", "Part 2.1 Q3, Part 3 Q8", ""),

    # Dynamic Mapping
    ("Dynamic field mapping (CSV / XLSX -> Google Sheets)", "Part 2.4 Q1, Q3", ""),
    ("Automated population with duplicate detection (content + filename)",
     "Part 2.4 Q5; Part 2.8 Q3", ""),

    # Reports
    ("One-page reports per uploaded source files", "Part 2.6 (all items)", ""),
    ("ABC Analysis from WMS data with visualization", "Part 2.5 (all items)",
     "Visualization = exported Google Sheet tabs"),

    # Agents
    ("Conversation Agent (NL interpreter)", "Part 2.3 (implicit in all items)",
     "Tester observes that NL requests are correctly interpreted"),
    ("Supervisor Agent risk evaluation", "Part 2.3 Q6, Part 2.9 Q5, Part 3 Q3", ""),

    # Gmail
    ("Search emails (date filters, sender, subject, attachments)",
     "Part 2.3 Q1", ""),
    ("Draft, send, reply, forward with thread preservation", "Part 2.3 Q1", ""),
    ("Manage email labels", "Part 2.3 Q1",
     "Folded into the Gmail-tasks question"),

    # Drive / Sheets / Docs
    ("Google Sheets API for data sync", "Part 2.3 Q4", ""),
    ("Google Drive: upload, create folders, list, search", "Part 2.3 Q3", ""),
    ("Google Docs templates / generation", "Part 2.3 Q4", ""),

    # Orchestration
    ("Multi-step workflow orchestration", "Part 2.3 Q5", ""),
    ("Pending task review interface (approve / reject)", "Part 2.9 Q5", ""),
    ("Risk classification: Safe / Moderate / Dangerous", "Part 2.3 Q6, Part 3 Q3",
     "User-observable as 'approval required for dangerous actions'"),

    # Audit / Security
    ("Audit trails / system logging", "Part 2.9 Q3", ""),
    ("JWT auth + OAuth 2.0 (Google Workspace)", "Part 2.1 Q1",
     "User-observable as 'sign-in worked'"),
    ("Three roles: Administrator / Manager / User (RBAC)",
     "Part 2.1 Q2, Part 3 Q5", ""),

    # UI / UX
    ("Document upload UI with PDF preview, page navigation, side-by-side "
     "highlight overlay for chunk locations",
     "Part 2.8 Q2", "NEW item added in v3 to plug this gap"),
    ("Comprehensive logs (LLM calls, agent execution, request summaries, "
     "pending action history)",
     "Part 2.9 Q3", ""),
    ("Real-time progress updates via WebSocket (current agent, step, "
     "completion)",
     "Part 2.3 Q9 (partial)",
     "Frontend displays stage / step / counter / elapsed time. The "
     "explicit 'which agent is processing' label is sent by the backend "
     "but not currently rendered in the InlineChatProgress UI; testers "
     "infer the agent from the tool name. Documented as a known gap."),
    ("Rate limiting / debouncing", "Part 3 Q7, Q8",
     "User-observable as 'no rate-limit surprises'"),

    # Architectural / not user-testable
    ("CORS / security headers", "(not user-testable)",
     "Architectural - validated via developer-side code review"),
    ("OAuth 2.0 with HS256 / RS256 token decoding",
     "Part 2.1 Q1 (sign-in works)",
     "Token-format detail not user-visible; covered by sign-in success"),
    ("User-specific session management (data isolation)",
     "Part 2.2 Q5, Part 3 Q5", ""),
]

st = doc.add_table(rows=len(scope_rows), cols=3)
st.style = "Light Grid Accent 1"
for i, (a, b, c) in enumerate(scope_rows):
    st.rows[i].cells[0].text = a
    st.rows[i].cells[1].text = b
    st.rows[i].cells[2].text = c
    if i == 0:
        for cc in st.rows[i].cells:
            shade_cell(cc, "DCE6F1")
            cc.paragraphs[0].runs[0].bold = True
set_col_widths(st, (8.5, 4.0, 5.0))

doc.add_page_break()

# ============================================================
# APPENDIX C - REPORTING TEMPLATE
# ============================================================
add_h("Appendix C - Reporting template for Chapter 4", level=1)
add_para(
    "When reporting results, fill in this template. The numbers below "
    "directly feed Chapter 4's 'Test Case and User Acceptance Test "
    "Analysis' section, replacing the '___%' placeholders.",
    italic=True,
)
add_blank()

add_h("C.1  Demographics", level=2)
add_para(
    "  - Total testers (N) = ___ (Administrator: ___, Manager: ___, User: ___).\n"
    "  - Mean hands-on session duration = ___ minutes."
)

add_h("C.2  System Usability Scale (SUS)", level=2)
add_para(
    "  - Mean SUS score = ___ / 100   (interpretation: ___).\n"
    "  - Per-item means + standard deviations are tabulated in Table ___ "
    "(append to Chapter 4)."
)

add_h("C.3  Functional acceptance - per section Yes-rate", level=2)
add_para(
    "Report the percentage of testers who answered Yes per item, "
    "averaged across all items in the section.\n"
    "  - 2.1 Authentication, Roles, Access:    Yes-rate = ___ %  (N applicable = ___)\n"
    "  - 2.2 SFX Bot:                          Yes-rate = ___ %  (N applicable = ___)\n"
    "  - 2.3 AI Assistant (Admin + Manager):   Yes-rate = ___ %  (N applicable = ___)\n"
    "  - 2.4 Dynamic Mapping:                  Yes-rate = ___ %  (N applicable = ___)\n"
    "  - 2.5 ABC Analysis:                     Yes-rate = ___ %  (N applicable = ___)\n"
    "  - 2.6 OPR:                              Yes-rate = ___ %  (N applicable = ___)\n"
    "  - 2.7 Workload Analysis:                Yes-rate = ___ %  (N applicable = ___)\n"
    "  - 2.8 KB Management (Admin only):       Yes-rate = ___ %  (N applicable = ___)\n"
    "  - 2.9 Account / Activity (Admin only):  Yes-rate = ___ %  (N applicable = ___)\n"
    "  - Highest-scoring section: ___ (Yes-rate = ___ %).\n"
    "  - Lowest-scoring section:  ___ (Yes-rate = ___ %)."
)

add_h("C.4  Quality scale (Likert)", level=2)
add_para(
    "  - Mean across 8 items = ___ / 5.\n"
    "  - Highest-scoring item: Q__ (mean = ___ ).\n"
    "  - Lowest-scoring item:  Q__ (mean = ___ )."
)

add_h("C.5  Adoption scale (Likert)", level=2)
add_para(
    "  - Mean across 6 items = ___ / 5.\n"
    "  - 'Would adopt tomorrow' (B3): mean = ___ , % who scored 4 or 5 = ___ %.\n"
    "  - 'Would recommend' (B5): mean = ___ , % who scored 4 or 5 = ___ %."
)

add_h("C.6  Open-ended thematic analysis", level=2)
add_para(
    "  - Top 3 recurring positive themes: ___, ___, ___.\n"
    "  - Top 3 recurring frustration themes: ___, ___, ___.\n"
    "  - Most-requested feature improvement: ___ (cited by ___ testers)."
)

add_h("C.7  Verdict distribution", level=2)
add_para(
    "  - Accept: ___ testers (___%).\n"
    "  - Accept with Conditions: ___ testers (___%).\n"
    "  - Reject: ___ testers (___%)."
)

add_h("C.8  Overall acceptance decision", level=2)
add_para(
    "Per the criteria stated on the cover page, the overall verdict is:  "
    f"{CHECKBOX} PASS    {CHECKBOX} CONDITIONAL PASS    {CHECKBOX} REJECT"
)
add_blank()

add_para(
    "End of UAT-FeatureScale.docx. This is the recommended primary "
    "instrument for Chapter 4. Pair with UAT-Real.docx for a small "
    "group of deep testers if the panel asks for scenario-level evidence.",
    italic=True, size=9,
)

doc.save(OUT)
print(f"OK -> {OUT}")
