"""
document_format_extractor.py
Extract formatting and structure from Google Docs AND Word documents
Supports creating new documents from templates with placeholder replacement
"""

import os
from typing import Dict, List, Optional
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError


def get_google_service(service_name: str, version: str, credentials_dict: Dict):
    """Get Google API service from credentials passed by the supervisor."""
    creds = Credentials(
        token=credentials_dict["access_token"],
        refresh_token=credentials_dict.get("refresh_token"),
        token_uri=credentials_dict.get("token_uri", "https://oauth2.googleapis.com/token"),
        client_id=credentials_dict.get("client_id", ""),
        client_secret=credentials_dict.get("client_secret", ""),
    )
    if creds.refresh_token:
        try:
            from google.auth.transport.requests import Request
            creds.refresh(Request())
        except Exception:
            pass
    service = build(service_name, version, credentials=creds)
    return service


class DocumentFormatExtractor:
    """Extract formatting information from a reference document"""

    def __init__(self, credentials_dict: Dict):
        self.credentials_dict = credentials_dict
        self.docs_service = get_google_service("docs", "v1", credentials_dict)
        self.drive_service = get_google_service("drive", "v3", credentials_dict)

    def extract_document_structure(self, document_id: str) -> Dict:
        """
        Extract structure and formatting from a document.
        NOW SUPPORTS: Google Docs AND uploaded Word files
        
        Args:
            document_id: Google Drive file ID (can be Google Doc or .docx)
        
        Returns:
            Dict with document structure or error
        """
        try:
            from google.oauth2.credentials import Credentials
            from googleapiclient.discovery import build
            
            # Build credentials
            creds = Credentials(
                token=self.credentials_dict.get("access_token"),
                refresh_token=self.credentials_dict.get("refresh_token"),
                token_uri=self.credentials_dict.get("token_uri", "https://oauth2.googleapis.com/token"),
                client_id=self.credentials_dict.get("client_id"),
                client_secret=self.credentials_dict.get("client_secret")
            )
            
            # First, check what type of file this is using Drive API
            drive_service = build('drive', 'v3', credentials=creds)
            
            try:
                file_metadata = drive_service.files().get(
                    fileId=document_id, 
                    fields='name, mimeType'
                ).execute()
                
                mime_type = file_metadata.get('mimeType')
                file_name = file_metadata.get('name')
                
                print(f" File: {file_name}")
                print(f" MIME Type: {mime_type}")
                
            except Exception as e:
                return {"error": f"Cannot access file: {str(e)}"}
            
            # Handle based on file type
            if mime_type == 'application/vnd.google-apps.document':
                # Native Google Doc - use Docs API
                print("   Type: Google Docs (native) → Using Docs API")
                return self._extract_google_doc_structure(document_id, creds)
            
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Word .docx file - extract text and treat as template
                print("   Type: Word .docx → Extracting as template")
                return self._extract_docx_structure(document_id, drive_service, file_name)
            
            elif mime_type == 'application/msword':
                # Legacy .doc file
                print("   Type: Legacy Word .doc → Converting via Drive")
                try:
                    text_content = drive_service.files().export(
                        fileId=document_id, 
                        mimeType='text/plain'
                    ).execute().decode('utf-8')
                    
                    return {
                        "title": file_name,
                        "content_blocks": [{"text": text_content, "type": "paragraph"}],
                        "raw_text": text_content,
                        "file_type": "doc"
                    }
                except:
                    return {"error": "Cannot read legacy .doc format. Please convert to .docx or Google Docs."}
            
            else:
                return {
                    "error": f"Unsupported file type: {mime_type}\n"
                            f"Supported types:\n"
                            f"- Google Docs (application/vnd.google-apps.document)\n"
                            f"- Word .docx (application/vnd.openxmlformats-officedocument.wordprocessingml.document)\n"
                            f"\nPlease convert your file to one of these formats."
                }
            
        except Exception as error:
            import traceback
            traceback.print_exc()
            return {"error": f"Error extracting document: {error}"}

    def _extract_google_doc_structure(self, document_id: str, creds) -> Dict:
        """
        Extract structure from a native Google Doc using Docs API
        """
        try:
            from googleapiclient.discovery import build
            
            docs_service = build('docs', 'v1', credentials=creds)
            document = docs_service.documents().get(documentId=document_id).execute()
            
            # Extract title
            title = document.get('title', 'Untitled')
            
            # Extract content blocks
            content_blocks = []
            body = document.get('body', {})
            content = body.get('content', [])
            
            for element in content:
                if 'paragraph' in element:
                    paragraph = element['paragraph']
                    text_content = ""
                    elements = []
                    
                    for text_element in paragraph.get('elements', []):
                        if 'textRun' in text_element:
                            text_run = text_element['textRun']
                            text = text_run.get('content', '')
                            text_content += text
                            
                            # Store element with style info
                            text_style = text_run.get('textStyle', {})
                            elements.append({
                                'text': text,
                                'style': {
                                    'bold': text_style.get('bold', False),
                                    'italic': text_style.get('italic', False),
                                    'underline': text_style.get('underline', False),
                                    'font_family': text_style.get('weightedFontFamily', {}).get('fontFamily', 'Arial'),
                                    'font_size': text_style.get('fontSize', {}).get('magnitude', 11),
                                    'foreground_color': text_style.get('foregroundColor', {}),
                                }
                            })
                    
                    if text_content.strip():
                        paragraph_style = paragraph.get('paragraphStyle', {})
                        content_blocks.append({
                            'type': 'paragraph',
                            'text': text_content,
                            'elements': elements,
                            'style': {
                                'alignment': paragraph_style.get('alignment', 'START'),
                                'heading': paragraph_style.get('namedStyleType', 'NORMAL_TEXT'),
                                'indent_start': paragraph_style.get('indentStart', {}).get('magnitude', 0),
                                'line_spacing': paragraph_style.get('lineSpacing', 100)
                            }
                        })
                
                elif 'table' in element:
                    table = element['table']
                    content_blocks.append({
                        'type': 'table',
                        'rows': table.get('rows', 0),
                        'columns': table.get('columns', 0)
                    })
            
            return {
                "title": title,
                "document_id": document_id,
                "content_blocks": content_blocks,
                "file_type": "google_docs"
            }
            
        except Exception as e:
            return {"error": f"Error reading Google Doc: {str(e)}"}

    def _extract_docx_structure(self, document_id: str, drive_service, file_name: str) -> Dict:
        """
        Extract structure from a Word .docx file
        """
        try:
            import io
            from docx import Document
            from googleapiclient.http import MediaIoBaseDownload
            
            # Download the .docx file
            request = drive_service.files().get_media(fileId=document_id)
            file_content = io.BytesIO()
            downloader = MediaIoBaseDownload(file_content, request)
            
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            file_content.seek(0)
            
            # Parse with python-docx
            doc = Document(file_content)
            
            # Extract content blocks
            content_blocks = []
            
            # Extract paragraphs with formatting
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    style_name = paragraph.style.name if paragraph.style else "Normal"
                    
                    content_blocks.append({
                        'type': 'paragraph',
                        'text': paragraph.text,
                        'style': style_name,
                        'elements': []  # Simplified for .docx
                    })
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                content_blocks.append({
                    'type': 'table',
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'data': table_data
                })
            
            print(f" Extracted {len(content_blocks)} content blocks from .docx")
            
            return {
                "title": file_name.replace('.docx', ''),
                "document_id": document_id,
                "content_blocks": content_blocks,
                "file_type": "docx"
            }
            
        except ImportError:
            return {
                "error": "python-docx library not installed. Install with: pip install python-docx"
            }
        except Exception as e:
            import traceback
            traceback.print_exc()
            return {"error": f"Error reading .docx file: {str(e)}"}

    def identify_placeholders(self, structure: Dict) -> List[str]:
        """
        Identify placeholder variables in the document structure.
        Works with both Google Docs and Word documents.
        
        Looks for patterns like:
        - [PLACEHOLDER]
        - {{PLACEHOLDER}}
        - {PLACEHOLDER}
        - <<PLACEHOLDER>>
        - key:____ (blank line format)
        
        Args:
            structure: Document structure from extract_document_structure
        
        Returns:
            List of unique placeholder names
        """
        import re
        
        if "error" in structure:
            return []
        
        placeholders = set()
        
        # Placeholder patterns to search for
        patterns = [
            r'\[([A-Z_][A-Z0-9_]*)\]',           # [PLACEHOLDER]
            r'\{\{([A-Z_][A-Z0-9_]*)\}\}',       # {{PLACEHOLDER}}
            r'\{([A-Z_][A-Z0-9_]*)\}',           # {PLACEHOLDER}
            r'<<([A-Z_][A-Z0-9_]*)>>',           # <<PLACEHOLDER>>
            r'\$\{([A-Z_][A-Z0-9_]*)\}',         # ${PLACEHOLDER}
        ]
        
        # Pattern for blank line placeholders
        blank_line_pattern = r'([a-zA-Z\s]+):\s*_{2,}'
        
        # Extract text from content blocks
        content_blocks = structure.get("content_blocks", [])
        
        for block in content_blocks:
            text = block.get("text", "")
            
            # Check against all patterns
            for pattern in patterns:
                matches = re.findall(pattern, text)
                placeholders.update(matches)
            
            # Check for blank line placeholders
            blank_matches = re.findall(blank_line_pattern, text, re.IGNORECASE)
            for match in blank_matches:
                normalized = match.strip().upper().replace(' ', '_')
                placeholders.add(normalized)
                print(f" Detected blank line placeholder: '{match}:____' → [{normalized}]")
            
            # Also check table data if present
            if block.get("type") == "table" and "data" in block:
                for row in block["data"]:
                    for cell in row:
                        for pattern in patterns:
                            matches = re.findall(pattern, cell)
                            placeholders.update(matches)
                        
                        # Check for blank lines in tables
                        blank_matches = re.findall(blank_line_pattern, cell, re.IGNORECASE)
                        for match in blank_matches:
                            normalized = match.strip().upper().replace(' ', '_')
                            placeholders.add(normalized)
        
        # Also check raw_text if available
        raw_text = structure.get("raw_text", "")
        if raw_text:
            for pattern in patterns:
                matches = re.findall(pattern, raw_text)
                placeholders.update(matches)
            
            blank_matches = re.findall(blank_line_pattern, raw_text, re.IGNORECASE)
            for match in blank_matches:
                normalized = match.strip().upper().replace(' ', '_')
                placeholders.add(normalized)
        
        result = sorted(list(placeholders))
        print(f" Total placeholders found: {result}")
        
        return result

    def create_from_template(
        self,
        template_document_id: str,
        new_title: str,
        placeholder_values: Dict[str, str] = None
    ) -> Dict:
        """
        Create a new document from a template (Google Doc or Word file)
        
        Args:
            template_document_id: ID of template (Google Doc or .docx in Drive)
            new_title: Title for the new document
            placeholder_values: Dict of placeholder replacements
        
        Returns:
            Dict with success status and new document details
        """
        try:
            from google.oauth2.credentials import Credentials
            from googleapiclient.discovery import build
            import re
            
            if placeholder_values is None:
                placeholder_values = {}
            
            # Build credentials
            creds = Credentials(
                token=self.credentials_dict.get("access_token"),
                refresh_token=self.credentials_dict.get("refresh_token"),
                token_uri=self.credentials_dict.get("token_uri", "https://oauth2.googleapis.com/token"),
                client_id=self.credentials_dict.get("client_id"),
                client_secret=self.credentials_dict.get("client_secret")
            )
            
            drive_service = build('drive', 'v3', credentials=creds)
            
            # Check file type
            file_metadata = drive_service.files().get(
                fileId=template_document_id,
                fields='name, mimeType'
            ).execute()
            
            mime_type = file_metadata.get('mimeType')
            file_name = file_metadata.get('name', 'Template')
            
            print(f" Template: {file_name} ({mime_type})")
            
            # Handle based on file type
            if mime_type == 'application/vnd.google-apps.document':
                # Native Google Doc - use copy method
                print("   Using Google Docs copy method")
                return self._create_from_google_doc(
                    template_document_id, new_title, placeholder_values, drive_service, creds
                )
            
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Word .docx - convert to Google Doc first
                print("   Converting Word .docx to Google Doc")
                return self._create_from_docx(
                    template_document_id, new_title, placeholder_values, drive_service, creds
                )
            
            else:
                return {
                    "success": False,
                    "error": f"Unsupported template type: {mime_type}"
                }
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": f"Error creating document: {str(e)}"
            }

    def _create_from_google_doc(
        self,
        template_document_id: str,
        new_title: str,
        placeholder_values: Dict[str, str],
        drive_service,
        creds
    ) -> Dict:
        """Create from a native Google Doc template"""
        try:
            # Step 1: Copy the template
            copied_file = drive_service.files().copy(
                fileId=template_document_id,
                body={'name': new_title}
            ).execute()
            
            new_doc_id = copied_file['id']
            print(f" Created copy: {new_doc_id}")
            
            # Step 2: Replace placeholders
            if placeholder_values:
                from googleapiclient.discovery import build
                docs_service = build('docs', 'v1', credentials=creds)
                
                requests = []
                for placeholder, value in placeholder_values.items():
                    # Try multiple placeholder formats
                    patterns = [
                        f"[{placeholder}]",
                        f"{{{{{placeholder}}}}}",  # {{PLACEHOLDER}}
                        f"{{{placeholder}}}",      # {PLACEHOLDER}
                        f"<<{placeholder}>>",
                        f"${{{placeholder}}}",
                        f"{placeholder.lower()}:____",  # blank line format
                        f"{placeholder.replace('_', ' ').title()}:____"  # "DATE" → "Date:____"
                    ]
                    
                    for pattern in patterns:
                        requests.append({
                            'replaceAllText': {
                                'containsText': {
                                    'text': pattern,
                                    'matchCase': False
                                },
                                'replaceText': str(value)
                            }
                        })
                
                if requests:
                    docs_service.documents().batchUpdate(
                        documentId=new_doc_id,
                        body={'requests': requests}
                    ).execute()
                    print(f" Replaced {len(placeholder_values)} placeholders")
            
            doc_url = f"https://docs.google.com/document/d/{new_doc_id}/edit"
            
            return {
                "success": True,
                "document_id": new_doc_id,
                "url": doc_url,
                "title": new_title,
                "template_used": template_document_id
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error creating from Google Doc: {str(e)}"
            }

    def _create_from_docx(
        self,
        template_document_id: str,
        new_title: str,
        placeholder_values: Dict[str, str],
        drive_service,
        creds
    ) -> Dict:
        """Create from a Word .docx template by converting to Google Doc"""
        try:
            import io
            from docx import Document
            from googleapiclient.http import MediaIoBaseDownload, MediaInMemoryUpload
            import re
            
            print(" Downloading Word template...")
            
            # Step 1: Download the .docx file
            request = drive_service.files().get_media(fileId=template_document_id)
            file_content = io.BytesIO()
            downloader = MediaIoBaseDownload(file_content, request)
            
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            file_content.seek(0)
            
            # Step 2: Parse with python-docx and replace placeholders
            doc = Document(file_content)
            
            if placeholder_values:
                print(f" Replacing {len(placeholder_values)} placeholders...")
                
                # Replace in paragraphs
                for paragraph in doc.paragraphs:
                    for placeholder, value in placeholder_values.items():
                        patterns = [
                            f"[{placeholder}]",
                            f"{{{{{placeholder}}}}}",
                            f"{{{placeholder}}}",
                            f"<<{placeholder}>>",
                            f"${{{placeholder}}}",
                            f"{placeholder.lower()}:____",
                            f"{placeholder.replace('_', ' ').title()}:____"
                        ]
                        
                        for pattern in patterns:
                            if pattern in paragraph.text:
                                # Replace in runs to preserve formatting
                                for run in paragraph.runs:
                                    if pattern in run.text:
                                        run.text = run.text.replace(pattern, str(value))
                
                # Replace in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for placeholder, value in placeholder_values.items():
                                    patterns = [
                                        f"[{placeholder}]",
                                        f"{{{{{placeholder}}}}}",
                                        f"{{{placeholder}}}",
                                        f"<<{placeholder}>>",
                                        f"${{{placeholder}}}",
                                        f"{placeholder.lower()}:____",
                                        f"{placeholder.replace('_', ' ').title()}:____"
                                    ]
                                    
                                    for pattern in patterns:
                                        if pattern in paragraph.text:
                                            for run in paragraph.runs:
                                                if pattern in run.text:
                                                    run.text = run.text.replace(pattern, str(value))
            
            # Step 3: Save modified .docx to memory
            output_docx = io.BytesIO()
            doc.save(output_docx)
            output_docx.seek(0)
            
            print(" Uploading modified document to Drive...")
            
            # Step 4: Upload as Google Doc (auto-converts)
            file_metadata = {
                'name': new_title,
                'mimeType': 'application/vnd.google-apps.document'
            }
            
            media = MediaInMemoryUpload(
                output_docx.read(),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                resumable=True
            )
            
            uploaded_file = drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id, webViewLink'
            ).execute()
            
            new_doc_id = uploaded_file['id']
            doc_url = uploaded_file['webViewLink']
            
            print(f" Created Google Doc: {new_doc_id}")
            
            return {
                "success": True,
                "document_id": new_doc_id,
                "url": doc_url,
                "title": new_title,
                "template_used": template_document_id,
                "converted_from": "docx"
            }
            
        except ImportError:
            return {
                "success": False,
                "error": "python-docx library not installed. Install with: pip install python-docx"
            }
        except Exception as e:
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": f"Error creating from .docx: {str(e)}"
            }

    def list_my_docs(self, query: str = "") -> List[Dict]:
        """
        List Google Docs AND Word documents in user's Drive
        
        Args:
            query: Optional search query
        
        Returns:
            List of documents with id, name, url, and type
        """
        try:
            # Search for both Google Docs and Word documents
            mime_types = [
                "application/vnd.google-apps.document",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ]
            
            search_query = "(" + " or ".join([f"mimeType='{mt}'" for mt in mime_types]) + ")"
            
            if query:
                search_query += f" and name contains '{query}'"
            
            results = (
                self.drive_service.files()
                .list(
                    q=search_query,
                    pageSize=20,
                    fields="files(id, name, mimeType, modifiedTime, webViewLink)",
                )
                .execute()
            )
            
            files = results.get("files", [])
            
            docs = []
            for file in files:
                doc_type = "Google Doc" if file['mimeType'] == 'application/vnd.google-apps.document' else "Word (.docx)"
                
                docs.append(
                    {
                        "id": file["id"],
                        "name": file["name"],
                        "url": file["webViewLink"],
                        "modified": file["modifiedTime"],
                        "type": doc_type,
                        "mime_type": file["mimeType"]
                    }
                )
            
            return docs
            
        except HttpError as error:
            return []


# Example usage
if __name__ == "__main__":
    from dotenv import load_dotenv
    
    load_dotenv()
    
    # Test credentials
    test_creds = {
        "access_token": os.getenv("GOOGLE_ACCESS_TOKEN"),
        "refresh_token": os.getenv("GOOGLE_REFRESH_TOKEN"),
    }
    
    extractor = DocumentFormatExtractor(test_creds)
    
    # Example: List user's documents (both Google Docs and Word files)
    print(" Your Documents:")
    docs = extractor.list_my_docs()
    for i, doc in enumerate(docs[:5], 1):
        print(f"{i}. [{doc['type']}] {doc['name']} (ID: {doc['id']})")
    
    # Example: Extract format from a template
    # template_id = "YOUR_TEMPLATE_DOC_ID"
    # structure = extractor.extract_document_structure(template_id)
    # print("\n Template Structure:")
    # print(f"Title: {structure['title']}")
    # print(f"File Type: {structure['file_type']}")
    # print(f"Blocks: {len(structure['content_blocks'])}")
    
    # Example: Identify placeholders
    # placeholders = extractor.identify_placeholders(structure)
    # print(f"\n Placeholders: {placeholders}")
    
    # Example: Create from template
    # result = extractor.create_from_template(
    #     template_document_id="TEMPLATE_ID",
    #     new_title="Team Meeting - Jan 15, 2025",
    #     placeholder_values={
    #         "DATE": "January 15, 2025",
    #         "TIME": "2:00 PM",
    #         "VENUE": "Conference Room A"
    #     }
    # )
    # print(f"\n Created: {result['url']}")