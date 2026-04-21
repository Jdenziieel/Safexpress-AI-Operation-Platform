#GOOGLE DOCS TOOLS
import os
import re
import json
from typing import Dict, Any
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from document_format_extractor import DocumentFormatExtractor
from typing import Dict, Any, Optional


_DOC_URL_RE = re.compile(r"/document/d/([A-Za-z0-9_-]+)")
_FILE_URL_RE = re.compile(r"/file/d/([A-Za-z0-9_-]+)")


def _extract_document_id(doc_id_or_url: str) -> str:
    """Extract a Docs documentId from a full Google URL or return the input as-is.

    Accepts the forms:
      - raw documentId, e.g. "1a2B3cD...XYZ"
      - https://docs.google.com/document/d/<id>/edit
      - https://drive.google.com/file/d/<id>/view
    The helper is a no-op for strings that already look like a clean ID, so it
    is always safe to call at the top of a tool implementation.
    """
    if not doc_id_or_url:
        return doc_id_or_url
    match = _DOC_URL_RE.search(doc_id_or_url) or _FILE_URL_RE.search(doc_id_or_url)
    if match:
        return match.group(1)
    return doc_id_or_url.strip()


def get_google_service(service_name: str, version: str, credentials_dict: Dict):
    """Get Google service with proper credential handling"""
    
    # Extract client_id and client_secret FIRST
    client_id = credentials_dict.get("client_id")
    client_secret = credentials_dict.get("client_secret")
    
    # If not provided in request, load from credentials.json
    if not client_id or not client_secret:
        try:
            # Look in multiple locations
            creds_paths = [
                os.path.join(os.path.dirname(__file__), 'key', 'credentials.json'),
                'key/credentials.json',
                'credentials.json'
            ]
            
            creds_file = None
            for path in creds_paths:
                if os.path.exists(path):
                    print(f" Loading credentials from: {path}")
                    with open(path, 'r') as f:
                        creds_file = json.load(f)
                    break
            
            if creds_file:
                if 'installed' in creds_file:
                    client_id = creds_file['installed']['client_id']
                    client_secret = creds_file['installed']['client_secret']
                elif 'web' in creds_file:
                    client_id = creds_file['web']['client_id']
                    client_secret = creds_file['web']['client_secret']
                
                print(f" Loaded client_id: {client_id[:20]}...")
            else:
                raise FileNotFoundError("credentials.json not found in any expected location")
                
        except Exception as e:
            print(f" Failed to load credentials.json: {e}")
            raise Exception(f"Missing GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET. Add them to .env or credentials.json")
    
    creds = Credentials(
        token=credentials_dict.get("access_token"),
        refresh_token=credentials_dict.get("refresh_token"),
        token_uri="https://oauth2.googleapis.com/token",
        client_id=client_id,
        client_secret=client_secret,
    )
    
    if creds.refresh_token:
        try:
            from google.auth.transport.requests import Request
            creds.refresh(Request())
        except Exception:
            pass
    
    print(f" Building {service_name} service with:")
    print(f"  - token: {'present' if creds.token else 'MISSING'}")
    print(f"  - refresh_token: {'present' if creds.refresh_token else 'MISSING'}")
    print(f"  - client_id: {'present' if client_id else 'MISSING'}")
    print(f"  - client_secret: {'present' if client_secret else 'MISSING'}")
    
    service = build(service_name, version, credentials=creds)
    
    return service


# Plain function without decorator - can be used directly or wrapped
def _create_google_doc_impl(
    title: str,
    credentials_dict: Dict,
    folder_id: str = None,
) -> str:
    """Create a new Google Doc, optionally moved into a target Drive folder.

    folder_id: Drive folder ID to place the doc in. The Docs API does not
    accept a parent at creation time, so we create at the Drive root and
    then reparent via the Drive API. If the move fails, the doc still exists
    at root and the failure is surfaced via a 'folder_moved: no' marker in
    the returned string (parser picks it up for the structured response).
    """
    try:
        docs_service = get_google_service("docs", "v1", credentials_dict)
        doc = {"title": title}

        document = docs_service.documents().create(body=doc).execute()

        doc_id = document.get("documentId")
        doc_url = f"https://docs.google.com/document/d/{doc_id}/edit"

        folder_moved = False
        folder_error = None
        if folder_id:
            try:
                drive_service = get_google_service("drive", "v3", credentials_dict)
                current = (
                    drive_service.files()
                    .get(fileId=doc_id, fields="parents")
                    .execute()
                )
                current_parents = ",".join(current.get("parents") or [])
                drive_service.files().update(
                    fileId=doc_id,
                    addParents=folder_id,
                    removeParents=current_parents or None,
                    fields="id, parents",
                ).execute()
                folder_moved = True
            except Exception as move_err:
                folder_error = str(move_err)

        lines = [
            "Document created successfully!",
            f"Title: {title}",
            f"ID: {doc_id}",
            f"URL: {doc_url}",
        ]
        if folder_id:
            lines.append(f"Folder ID: {folder_id}")
            lines.append(f"Folder moved: {'yes' if folder_moved else 'no'}")
            if folder_error:
                lines.append(f"Folder move error: {folder_error}")
        return "\n".join(lines)

    except HttpError as error:
        return f"error in creating document: {error}"
    except KeyError as error:
        return f"Missing credentials: {error}"
    except Exception as error:
        return f"Unexpected error: {error}"


def _add_text_to_doc_impl(document_id: str, text: str, credentials_dict: Dict) -> str:
    """Implementation of adding text to a Google Doc"""
    try:
        document_id = _extract_document_id(document_id)
        docs_service = get_google_service("docs", "v1", credentials_dict)
        # Create the request body for inserting text
        requests = [{"insertText": {"location": {"index": 1}, "text": text}}]
        result = (
            docs_service.documents()
            .batchUpdate(documentId=document_id, body={"requests": requests})
            .execute()
        )
        doc_url = f"https://docs.google.com/document/d/{document_id}/edit"
        return f"Text added successfully!\nDocument ID: {document_id}\nURL: {doc_url}"
    except HttpError as error:
        return f"error adding text to document: {error}"
    except KeyError as error:
        return f"Missing credentials: {error}"
    except Exception as error:
        return f"Unexpected error: {error}"


def _read_google_doc_impl(document_id: str, credentials_dict: Dict) -> str:
    """Implementation of reading text from a Google Doc"""
    try:
        document_id = _extract_document_id(document_id)
        docs_service = get_google_service("docs", "v1", credentials_dict)
        document = docs_service.documents().get(documentId=document_id).execute()
        # initilize empty string to collect text
        text = ""

        # gets the content array from document body
        content = document.get("body", {}).get("content", [])

        # loop through each element in the content
        for element in content:
            # this checks if the element is a paragraph
            if "paragraph" in element:
                # get the element in the paragraph
                paragraph_elements = element["paragraph"].get("elements", [])

                # loop through each element in the paragraph
                for para_element in paragraph_elements:
                    # check if the element contains text
                    if "textRun" in para_element:
                        # extract the text content and add it to the string
                        text += para_element["textRun"].get("content", "")

        # return the extracted text with document info
        doc_url = f"https://docs.google.com/document/d/{document_id}/edit"
        return (
            f"Document content:\n\n{text}\n\nDocument ID: {document_id}\nURL: {doc_url}"
        )
    except HttpError as error:
        return f"error reading document: {error}"
    except KeyError as error:
        return f"Missing field in document: {error}"
    except Exception as error:
        return f"Unexpected error: {error}"


def _edit_google_doc_impl(
    document_id: str,
    old_text: str = None,
    new_text: str = None,
    content: str = None,
    credentials_dict: Dict = None,
) -> str:
    """Implementation of editing/replacing text in a Google Doc

    Args:
        document_id: The ID of the document to edit
        old_text: The text to find and replace
        new_text: The replacement text
        content: Alias for new_text when the request uses generic content field
        credentials_dict: Google OAuth credentials

    Returns:
        Success message with edit details
    """
    if new_text is None and content is not None:
        new_text = content

    if not old_text or not new_text:
        return (
            "Error: edit_doc requires both old_text and new_text. "
            "Provide the exact text to replace and the replacement text, "
            "or use update_doc to replace the whole document content."
        )

    try:
        document_id = _extract_document_id(document_id)
        docs_service = get_google_service("docs", "v1", credentials_dict)

        # First, read the document to find the old text
        document = docs_service.documents().get(documentId=document_id).execute()

        # Extract full document text to find the position
        full_text = ""
        content = document.get("body", {}).get("content", [])

        for element in content:
            if "paragraph" in element:
                paragraph_elements = element["paragraph"].get("elements", [])
                for para_element in paragraph_elements:
                    if "textRun" in para_element:
                        full_text += para_element["textRun"].get("content", "")

        # Find the position of old_text
        if old_text not in full_text:
            return f"Error: Text '{old_text}' not found in document"

        # Calculate start and end indices
        start_index = full_text.index(old_text) + 1  # +1 because Docs API is 1-indexed
        end_index = start_index + len(old_text)

        # Create the batch update requests
        requests = [
            # First, delete the old text
            {
                "deleteContentRange": {
                    "range": {"startIndex": start_index, "endIndex": end_index}
                }
            },
            # Then, insert the new text at the same position
            {"insertText": {"location": {"index": start_index}, "text": new_text}},
        ]

        # Execute the batch update
        result = (
            docs_service.documents()
            .batchUpdate(documentId=document_id, body={"requests": requests})
            .execute()
        )

        doc_url = f"https://docs.google.com/document/d/{document_id}/edit"
        return (
            f"Text edited successfully!\n"
            f"Replaced: '{old_text}'\n"
            f"With: '{new_text}'\n"
            f"Document ID: {document_id}\n"
            f"URL: {doc_url}"
        )

    except HttpError as error:
        return f"Error editing document: {error}"
    except ValueError as error:
        return f"Text not found in document: {error}"
    except Exception as error:
        return f"Unexpected error: {error}"


def _update_entire_doc_impl(
    document_id: str, new_content: str, credentials_dict: Dict
) -> str:
    """Implementation of replacing entire document content

    Args:
        document_id: The ID of the document to update
        new_content: The new complete content for the document
        credentials_dict: Google OAuth credentials

    Returns:
        Success message with update details
    """
    try:
        document_id = _extract_document_id(document_id)
        docs_service = get_google_service("docs", "v1", credentials_dict)

        # First, get the document to find the end index
        document = docs_service.documents().get(documentId=document_id).execute()

        # Calculate the end index (total length of current content)
        end_index = document.get("body", {}).get("content", [{}])[-1].get("endIndex", 1)

        # Create batch update requests
        requests = []

        # Only delete if there's content to delete (end_index > 2)
        if end_index > 2:
            requests.append(
                {
                    "deleteContentRange": {
                        "range": {
                            "startIndex": 1,
                            "endIndex": end_index
                            - 1,  # -1 to avoid deleting the protected last character
                        }
                    }
                }
            )

        # Always insert new content at the beginning
        requests.append({"insertText": {"location": {"index": 1}, "text": new_content}})

        # Execute the batch update
        result = (
            docs_service.documents()
            .batchUpdate(documentId=document_id, body={"requests": requests})
            .execute()
        )

        doc_url = f"https://docs.google.com/document/d/{document_id}/edit"
        return (
            f"Document content updated successfully!\n"
            f"Document ID: {document_id}\n"
            f"New content length: {len(new_content)} characters\n"
            f"URL: {doc_url}"
        )

    except HttpError as error:
        return f"Error updating document: {error}"
    except Exception as error:
        return f"Unexpected error: {error}"


def _list_user_docs_impl(credentials_dict: Dict, search_query: str = "") -> str:
    """List user's Google Docs (to find their templates)

    Args:
        credentials_dict: Google OAuth credentials
        search_query: Optional search term (e.g., "template", "MOM")

    Returns:
        Formatted list of user's documents
    """
    try:
        extractor = DocumentFormatExtractor(credentials_dict)
        docs = extractor.list_my_docs(search_query)

        if not docs:
            return "No documents found. Please upload a template document to your Google Drive first."

        result = f"Found {len(docs)} document(s):\n\n"
        for i, doc in enumerate(docs, 1):
            result += f"{i}. {doc['name']}\n"
            result += f"   ID: {doc['id']}\n"
            result += f"   URL: {doc['url']}\n"
            result += f"   Modified: {doc['modified']}\n\n"

        return result

    except Exception as error:
        return f"Error listing documents: {error}"


def _extract_template_structure_impl(
    template_document_id: str, credentials_dict: Dict
) -> str:
    """Extract formatting and structure from a reference document

    Args:
        template_document_id: ID of the template document
        credentials_dict: Google OAuth credentials

    Returns:
        Summary of extracted structure
    """
    try:
        extractor = DocumentFormatExtractor(credentials_dict)
        structure = extractor.extract_document_structure(template_document_id)

        if "error" in structure:
            return structure["error"]

        # Identify placeholders
        placeholders = extractor.identify_placeholders(structure)

        result = f"Template structure extracted successfully.\n\n"
        result += f"Template: {structure['title']}\n"
        result += f"Content Blocks: {len(structure['content_blocks'])}\n\n"

        if placeholders:
            result += f"Found Placeholders:\n"
            for placeholder in placeholders:
                result += f"   - [{placeholder}]\n"
            result += f"\nYou can fill these when creating a new document.\n"
        else:
            result += "No placeholders found. This template will be copied as-is.\n"

        result += f"\nDocument ID: {template_document_id}"

        return result

    except Exception as error:
        return f"Error extracting template: {error}"


def _create_from_reference_impl(
    template_document_id: str,
    new_title: str,
    placeholder_values: Dict[str, str],
    credentials_dict: Dict,
) -> str:
    """Create a new document based on user's reference template

    Args:
        template_document_id: ID of the template document
        new_title: Title for the new document
        placeholder_values: Dict of placeholder replacements
            Example: {"DATE": "Jan 15, 2025", "VENUE": "Room A"}
        credentials_dict: Google OAuth credentials

    Returns:
        Success message with new document details
    """
    try:
        extractor = DocumentFormatExtractor(credentials_dict)

        result = extractor.create_from_template(
            template_document_id=template_document_id,
            new_title=new_title,
            placeholder_values=placeholder_values,
        )

        if result.get("success"):
            response = f"Document created from your template.\n\n"
            response += f"Title: {result['title']}\n"
            response += f"Template Used: {result['template_used']}\n"
            response += f"Document ID: {result['document_id']}\n"
            response += f"URL: {result['url']}\n"

            if placeholder_values:
                response += f"\nPlaceholders Filled:\n"
                for key, value in placeholder_values.items():
                    response += f"   - [{key}] -> {value}\n"

            return response
        else:
            return f"Error: {result.get('error')}"

    except Exception as error:
        return f"Error creating document: {error}"

def _create_from_uploaded_template_impl(
    template_file_id: str,
    new_title: str,
    placeholder_values: Optional[Dict[str, str]],
    credentials_dict: Dict,
    output_format: str = "google_docs"  # NEW: "google_docs" or "pdf"
) -> str:
    """
    Create document from uploaded template file in Drive
    
    Args:
        template_file_id: Google Drive file ID of template
        new_title: Title for new document
        placeholder_values: Dict of placeholder replacements
        credentials_dict: Google OAuth credentials
        output_format: "google_docs" (default, editable) or "pdf" (final output)
    
    Returns:
        Success message with document details
    """
    try: 
        print(f"\n{'='*60}")
        print(f" CREDENTIALS RECEIVED IN DOCS AGENT:")
        print(f"{'='*60}")
        print(f"access_token: {'present' if credentials_dict.get('access_token') else 'MISSING'}")
        print(f"refresh_token: {'present' if credentials_dict.get('refresh_token') else 'MISSING'}")
        print(f"client_id: {'present' if credentials_dict.get('client_id') else 'MISSING'}")
        print(f"client_secret: {'present' if credentials_dict.get('client_secret') else 'MISSING'}")
        print(f"{'='*60}\n")
        
        # Use existing document format extractor
        extractor = DocumentFormatExtractor(credentials_dict)
        
        # Convert placeholder_values from JSON string if needed
        if isinstance(placeholder_values, str):
            import json
            placeholder_values = json.loads(placeholder_values)
        
        # Create document from template (always creates Google Doc first)
        result = extractor.create_from_template(
            template_document_id=template_file_id,
            new_title=new_title,
            placeholder_values=placeholder_values or {}
        )
        
        if not result.get("success"):
            return f"Error: {result.get('error')}"
        
        # If user wants PDF output, export it
        if output_format == "pdf":
            print("Exporting document as PDF...")
            try:
                from googleapiclient.discovery import build
                from googleapiclient.http import MediaInMemoryUpload
                from google.oauth2.credentials import Credentials
                from google.auth.transport.requests import Request as AuthRequest
                
                creds = Credentials(
                    token=credentials_dict.get("access_token"),
                    refresh_token=credentials_dict.get("refresh_token"),
                    token_uri=credentials_dict.get("token_uri", "https://oauth2.googleapis.com/token"),
                    client_id=credentials_dict.get("client_id"),
                    client_secret=credentials_dict.get("client_secret")
                )
                if creds.refresh_token:
                    try:
                        creds.refresh(AuthRequest())
                    except Exception:
                        pass
                
                drive_service = build('drive', 'v3', credentials=creds)
                
                # Export Google Doc as PDF
                pdf_content = drive_service.files().export(
                    fileId=result['document_id'],
                    mimeType='application/pdf'
                ).execute()
                
                # Get parent folder of the Google Doc
                doc_metadata = drive_service.files().get(
                    fileId=result['document_id'],
                    fields='parents'
                ).execute()
                
                parents = doc_metadata.get('parents', [])
                
                # Upload PDF version to same folder
                pdf_metadata = {
                    'name': f"{new_title}.pdf",
                    'mimeType': 'application/pdf',
                    'parents': parents
                }
                
                pdf_file = drive_service.files().create(
                    body=pdf_metadata,
                    media_body=MediaInMemoryUpload(pdf_content, mimetype='application/pdf'),
                    fields='id, webViewLink'
                ).execute()
                
                # Build response for PDF output
                response = f"Document created from template and exported as PDF.\n\n"
                response += f"Title: {new_title}.pdf\n"
                response += f"PDF ID: {pdf_file['id']}\n"
                response += f"PDF URL: {pdf_file['webViewLink']}\n"
                response += f"Original Google Doc ID: {result['document_id']}\n"
                response += f"Google Doc URL: {result['url']}\n"
                
                if placeholder_values:
                    response += f"\nPlaceholders filled:\n"
                    for key, value in placeholder_values.items():
                        response += f"   - [{key}] -> {value}\n"
                
                response += f"\nFormat: PDF (non-editable, final output)"
                
                return response
                
            except Exception as pdf_error:
                print(f"PDF export failed: {pdf_error}")
                response = f"Created Google Doc but PDF export failed: {str(pdf_error)}\n\n"
                response += f"Google Doc Title: {result['title']}\n"
                response += f"Document ID: {result['document_id']}\n"
                response += f"URL: {result['url']}\n"
                return response
        
        response = f"Document created from uploaded template.\n\n"
        response += f"Title: {result['title']}\n"
        response += f"Document ID: {result['document_id']}\n"
        response += f"URL: {result['url']}\n"
        
        if placeholder_values:
            response += f"\nPlaceholders filled:\n"
            for key, value in placeholder_values.items():
                response += f"   - [{key}] -> {value}\n"
        
        response += f"\nFormat: Google Docs (editable)"
        
        return response
    
    except Exception as error:
        import traceback
        traceback.print_exc()
        return f"Error creating document from template: {error}"
    
def _analyze_uploaded_template_impl(
    template_file_id: str,
    credentials_dict: Dict
) -> str:
    """
    Analyze uploaded template to extract structure, placeholders, and formatting
    
    Args:
        template_file_id: Google Drive file ID of uploaded template
        credentials_dict: Google OAuth credentials
    
    Returns:
        JSON string with template analysis
    """
    try:


        print(f"\n{'='*60}")
        print(f" CREDENTIALS RECEIVED IN DOCS AGENT:")
        print(f"{'='*60}")
        print(f"access_token: {'present' if credentials_dict.get('access_token') else 'MISSING'}")
        print(f"refresh_token: {'present' if credentials_dict.get('refresh_token') else 'MISSING'}")
        print(f"client_id: {'present' if credentials_dict.get('client_id') else 'MISSING'}")
        print(f"client_secret: {'present' if credentials_dict.get('client_secret') else 'MISSING'}")
        print(f"{'='*60}\n")
        extractor = DocumentFormatExtractor(credentials_dict)
        
        # Extract structure
        structure = extractor.extract_document_structure(template_file_id)
        
        if "error" in structure:
            return json.dumps({
                "success": False,
                "error": structure["error"]
            })
        
        # Identify placeholders
        placeholders = extractor.identify_placeholders(structure)
        
        # Build analysis result
        analysis = {
            "success": True,
            "template_id": template_file_id,
            "title": structure.get("title", "Untitled"),
            "content_blocks": len(structure.get("content_blocks", [])),
            "placeholders": placeholders,
            "has_placeholders": len(placeholders) > 0,
            "structure_type": "structured" if placeholders else "unstructured",
            "ready_for_use": True
        }
        
        print(f"\n Template Analysis:")
        print(f"  Title: {analysis['title']}")
        print(f"  Blocks: {analysis['content_blocks']}")
        print(f"  Placeholders: {placeholders}")
        print(f"  Type: {analysis['structure_type']}")
        
        return json.dumps(analysis, indent=2)
        
    except Exception as error:
        import traceback
        traceback.print_exc()
        return json.dumps({
            "success": False,
            "error": str(error)
        })
def _create_from_existing_data_and_template_impl(
    template_file_name: str,
    data_file_name: str,
    new_title: str,
    credentials_dict: Dict,
    output_format: str = "google_docs"
) -> str:
    """
    Create document from existing template and data files in Google Drive (using file names)
    
    Args:
        template_file_name: Name of template file in Google Drive (e.g., "Monthly Report Template")
        data_file_name: Name of data file in Google Drive (e.g., "January Data.txt")
        new_title: Title for new document
        credentials_dict: Google OAuth credentials
        output_format: "google_docs" (default) or "pdf"
    
    Returns:
        Success message with document details
    """
    try:
        print(f"\n{'='*60}")
        print(f" CREATE FROM EXISTING DATA & TEMPLATE")
        print(f"{'='*60}")
        print(f"Template Name: {template_file_name}")
        print(f"Data Name: {data_file_name}")
        print(f"New Title: {new_title}")
        print(f"Output Format: {output_format}")
        print(f"{'='*60}\n")
        
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request as AuthRequest
        from googleapiclient.discovery import build
        
        creds = Credentials(
            token=credentials_dict.get("access_token"),
            refresh_token=credentials_dict.get("refresh_token"),
            token_uri=credentials_dict.get("token_uri", "https://oauth2.googleapis.com/token"),
            client_id=credentials_dict.get("client_id"),
            client_secret=credentials_dict.get("client_secret")
        )
        if creds.refresh_token:
            try:
                creds.refresh(AuthRequest())
            except Exception:
                pass
        
        drive_service = build('drive', 'v3', credentials=creds)
        
        # Step 1: Search for template file by name
        print(f" Searching for template: '{template_file_name}'...")
        template_query = f"name='{template_file_name}' and trashed=false"
        template_results = drive_service.files().list(
            q=template_query,
            fields="files(id, name, mimeType)",
            pageSize=10
        ).execute()
        
        template_files = template_results.get('files', [])
        
        if not template_files:
            return f"Error: Template file '{template_file_name}' not found in Google Drive"
        
        # If multiple files with same name, prefer Google Docs
        template_file = None
        for file in template_files:
            if file['mimeType'] == 'application/vnd.google-apps.document':
                template_file = file
                break
        
        if not template_file:
            template_file = template_files[0]  # Use first match
        
        template_file_id = template_file['id']
        print(f" Found template: {template_file['name']} (ID: {template_file_id})")
        
        # Step 2: Search for data file by name
        print(f" Searching for data file: '{data_file_name}'...")
        data_query = f"name='{data_file_name}' and trashed=false"
        data_results = drive_service.files().list(
            q=data_query,
            fields="files(id, name, mimeType)",
            pageSize=10
        ).execute()
        
        data_files = data_results.get('files', [])
        
        if not data_files:
            return f"Error: Data file '{data_file_name}' not found in Google Drive"
        
        data_file = data_files[0]  # Use first match
        data_file_id = data_file['id']
        data_mime_type = data_file['mimeType']
        print(f" Found data file: {data_file['name']} (ID: {data_file_id}, Type: {data_mime_type})")
        
        # Step 3: Read data content based on type
        print(f" Reading data from '{data_file['name']}'...")
        
        if data_mime_type == 'application/vnd.google-apps.document':
            data_content = drive_service.files().export(fileId=data_file_id, mimeType='text/plain').execute().decode('utf-8')
        elif data_mime_type in ['text/plain', 'text/csv']:
            data_content = drive_service.files().get_media(fileId=data_file_id).execute().decode('utf-8')
        elif data_mime_type == 'application/vnd.google-apps.spreadsheet':
            data_content = drive_service.files().export(fileId=data_file_id, mimeType='text/csv').execute().decode('utf-8')
        else:
            return f"Error: Unsupported data file type: {data_mime_type}"
        
        print(f" Read {len(data_content)} characters from data file")
        
        # Step 4: Analyze template for placeholders
        extractor = DocumentFormatExtractor(credentials_dict)
        structure = extractor.extract_document_structure(template_file_id)
        
        if "error" in structure:
            return f"Error analyzing template: {structure['error']}"
        
        placeholders = extractor.identify_placeholders(structure)
        print(f" Found placeholders in template: {placeholders}")
        
        # Step 5: Parse data content into placeholder values
        placeholder_values = {}
        
        if not placeholders:
            # No placeholders, just use data as content
            placeholder_values = {"CONTENT": data_content}
            print("ℹ No placeholders found in template, using data as CONTENT")
        else:
            # Try to parse data as key-value pairs (flexible parsing)
            lines = data_content.strip().split('\n')
            
            for line in lines:
                # Support multiple separators: colon, equals, pipe
                separator = None
                if ':' in line:
                    separator = ':'
                elif '=' in line:
                    separator = '='
                elif '|' in line:
                    separator = '|'
                
                if separator:
                    parts = line.split(separator, 1)
                    if len(parts) == 2:
                        key = parts[0].strip().upper().replace(' ', '_').replace('-', '_')
                        value = parts[1].strip()
                        
                        # Match to placeholders (fuzzy matching)
                        matched = False
                        for placeholder in placeholders:
                            if key == placeholder or key in placeholder or placeholder in key:
                                placeholder_values[placeholder] = value
                                print(f" Matched: [{placeholder}] = {value[:50]}{'...' if len(value) > 50 else ''}")
                                matched = True
                                break
                        
                        # If not matched but looks like a valid key, store it
                        if not matched and key in placeholders:
                            placeholder_values[key] = value
            
            # Fill missing placeholders
            for placeholder in placeholders:
                if placeholder not in placeholder_values:
                    if placeholder == 'CONTENT' or placeholder == 'DATA' or placeholder == 'TEXT':
                        # Use entire data content for content-type placeholders
                        placeholder_values[placeholder] = data_content
                        print(f" ℹ Using full data for [{placeholder}]")
                    elif placeholder == 'DATE':
                        # Auto-fill date
                        from datetime import datetime
                        placeholder_values[placeholder] = datetime.now().strftime("%B %d, %Y")
                        print(f" ℹ Auto-filled [{placeholder}] with current date")
                    else:
                        # Leave empty
                        placeholder_values[placeholder] = ""
                        print(f" No data for [{placeholder}], leaving empty")
        
        print(f" Final placeholder mapping: {list(placeholder_values.keys())}")
        
        # Step 6: Create document from template
        print(f" Creating document '{new_title}' from template...")
        result = extractor.create_from_template(
            template_document_id=template_file_id,
            new_title=new_title,
            placeholder_values=placeholder_values
        )
        
        if not result.get("success"):
            return f"Error creating document: {result.get('error')}"
        
        print(f" Document created: {result['document_id']}")
        
        # NEW: Handle blank line replacements (date:____, attendees:____, etc.)
        print(f" Checking for blank line placeholders...")
        _replace_blank_lines_in_document(
            document_id=result['document_id'],
            placeholder_values=placeholder_values,
            credentials_dict=credentials_dict
        )
        
        # Step 7: Export as PDF if requested
        if output_format == "pdf":
            print(" Exporting as PDF...")
            try:
                pdf_content = drive_service.files().export(
                    fileId=result['document_id'],
                    mimeType='application/pdf'
                ).execute()
                
                doc_metadata = drive_service.files().get(
                    fileId=result['document_id'],
                    fields='parents'
                ).execute()
                
                parents = doc_metadata.get('parents', [])
                
                pdf_metadata = {
                    'name': f"{new_title}.pdf",
                    'mimeType': 'application/pdf',
                    'parents': parents
                }
                
                from googleapiclient.http import MediaInMemoryUpload
                pdf_file = drive_service.files().create(
                    body=pdf_metadata,
                    media_body=MediaInMemoryUpload(pdf_content, mimetype='application/pdf'),
                    fields='id, webViewLink'
                ).execute()
                
                response = f"Document created from existing files and exported as PDF.\n\n"
                response += f"Title: {new_title}.pdf\n"
                response += f"PDF ID: {pdf_file['id']}\n"
                response += f"PDF URL: {pdf_file['webViewLink']}\n"
                response += f"Google Doc ID: {result['document_id']}\n"
                response += f"Google Doc URL: {result['url']}\n"
                response += f"Data Source: {data_file_name}\n"
                response += f"Template Used: {template_file_name}\n"
                
                if placeholder_values:
                    response += f"\nPlaceholders filled: {len(placeholder_values)}\n"
                    for key in list(placeholder_values.keys())[:5]:
                        value = placeholder_values[key]
                        preview = value[:50] + "..." if len(value) > 50 else value
                        response += f"   - [{key}] -> {preview}\n"
                    if len(placeholder_values) > 5:
                        response += f"   ... and {len(placeholder_values) - 5} more\n"
                
                return response
                
            except Exception as pdf_error:
                print(f"PDF export failed: {pdf_error}")
                import traceback
                traceback.print_exc()
        
        response = f"Document created from existing files.\n\n"
        response += f"Title: {result['title']}\n"
        response += f"Document ID: {result['document_id']}\n"
        response += f"URL: {result['url']}\n"
        response += f"Data Source: {data_file_name}\n"
        response += f"Template Used: {template_file_name}\n"
        
        if placeholder_values:
            response += f"\nPlaceholders filled: {len(placeholder_values)}\n"
            for key in list(placeholder_values.keys())[:5]:
                value = placeholder_values[key]
                preview = value[:50] + "..." if len(value) > 50 else value
                response += f"   - [{key}] -> {preview}\n"
            if len(placeholder_values) > 5:
                response += f"   ... and {len(placeholder_values) - 5} more\n"
        
        return response
    
    except Exception as error:
        import traceback
        traceback.print_exc()
        return f"Error creating document from existing files: {error}"


def _replace_blank_lines_in_document(
    document_id: str,
    placeholder_values: Dict[str, str],
    credentials_dict: Dict
) -> None:
    """
    Replace blank line placeholders (date:____) in a Google Doc
    
    Args:
        document_id: Google Doc ID
        placeholder_values: Dict of {PLACEHOLDER: value}
        credentials_dict: OAuth credentials
    """
    try:
        docs_service = get_google_service("docs", "v1", credentials_dict)
        
        # Get current document content
        document = docs_service.documents().get(documentId=document_id).execute()
        
        # Extract full text
        full_text = ""
        content = document.get("body", {}).get("content", [])
        for element in content:
            if "paragraph" in element:
                for para_element in element["paragraph"].get("elements", []):
                    if "textRun" in para_element:
                        full_text += para_element["textRun"].get("content", "")
        
        # Build replacement requests
        requests = []
        
        for key, value in placeholder_values.items():
            # Try multiple variations of the key
            variations = [
                f"{key.lower()}:____",           # date:____
                f"{key.lower()}: ____",          # date: ____
                f"{key.replace('_', ' ').lower()}:____",  # company_name → company name:____
                f"{key.replace('_', ' ').title()}:____",  # company_name → Company Name:____
            ]
            
            for pattern in variations:
                if pattern in full_text.lower():
                    # Find exact position (case-insensitive search)
                    lower_text = full_text.lower()
                    start_pos = lower_text.index(pattern)
                    
                    # Get actual text from original (preserves case)
                    actual_pattern = full_text[start_pos:start_pos + len(pattern)]
                    
                    # Calculate indices (Google Docs API is 1-indexed)
                    start_index = start_pos + 1
                    end_index = start_index + len(actual_pattern)
                    
                    # Build replacement text (preserve original key format)
                    key_part = actual_pattern.split(':')[0]
                    replacement_text = f"{key_part}: {value}"
                    
                    # Add delete and insert requests
                    requests.append({
                        "deleteContentRange": {
                            "range": {
                                "startIndex": start_index,
                                "endIndex": end_index
                            }
                        }
                    })
                    requests.append({
                        "insertText": {
                            "location": {"index": start_index},
                            "text": replacement_text
                        }
                    })
                    
                    print(f" Replacing '{actual_pattern}' with '{replacement_text}'")
                    
                    # Update full_text for next iteration
                    full_text = full_text[:start_pos] + replacement_text + full_text[start_pos + len(actual_pattern):]
                    break  # Found match, move to next placeholder
        
        # Execute batch update if we have replacements
        if requests:
            print(f" Executing {len(requests)} replacement operations...")
            docs_service.documents().batchUpdate(
                documentId=document_id,
                body={"requests": requests}
            ).execute()
            print(" Blank line replacements completed")
        else:
            print("ℹ No blank line patterns found to replace")
    
    except Exception as e:
        print(f" Warning: Could not replace blank lines: {e}")
        # Don't fail the entire operation, just log warning

def _create_from_template_and_data_ids_impl(
    template_file_id: str,
    data_file_id: str,
    new_title: str,
    credentials_dict: Dict,
    output_format: str = "google_docs"
) -> str:
    """
    Create document from template and data files using their IDs
    (IDs already found by Drive Agent)
    
    Args:
        template_file_id: Google Drive file ID of template
        data_file_id: Google Drive file ID of data file
        new_title: Title for new document
        credentials_dict: Google OAuth credentials
        output_format: "google_docs" (default) or "pdf"
    
    Returns:
        Success message with document details
    """
    try:
        print(f"\n{'='*60}")
        print(f" CREATE FROM TEMPLATE & DATA (USING IDs)")
        print(f"{'='*60}")
        print(f"Template ID: {template_file_id}")
        print(f"Data ID: {data_file_id}")
        print(f"New Title: {new_title}")
        print(f"Output Format: {output_format}")
        print(f"{'='*60}\n")
        
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request as AuthRequest
        from googleapiclient.discovery import build
        
        creds = Credentials(
            token=credentials_dict.get("access_token"),
            refresh_token=credentials_dict.get("refresh_token"),
            token_uri=credentials_dict.get("token_uri", "https://oauth2.googleapis.com/token"),
            client_id=credentials_dict.get("client_id"),
            client_secret=credentials_dict.get("client_secret")
        )
        if creds.refresh_token:
            try:
                creds.refresh(AuthRequest())
            except Exception:
                pass
        
        drive_service = build('drive', 'v3', credentials=creds)
        
        # Step 1: Read data file content
        print(f" Reading data file...")
        data_file_metadata = drive_service.files().get(fileId=data_file_id, fields='name, mimeType').execute()
        data_mime_type = data_file_metadata['mimeType']
        data_file_name = data_file_metadata['name']
        
        print(f" Data file: {data_file_name} (Type: {data_mime_type})")
        
        # EXPANDED SUPPORT: Handle multiple file types
        if data_mime_type == 'application/vnd.google-apps.document':
            print("   Format: Google Docs → Exporting as plain text")
            data_content = drive_service.files().export(fileId=data_file_id, mimeType='text/plain').execute().decode('utf-8')
        
        elif data_mime_type in ['text/plain', 'text/csv']:
            print("   Format: Text/CSV → Reading directly")
            data_content = drive_service.files().get_media(fileId=data_file_id).execute().decode('utf-8')
        
        elif data_mime_type == 'application/vnd.google-apps.spreadsheet':
            print("   Format: Google Sheets → Exporting as CSV")
            data_content = drive_service.files().export(fileId=data_file_id, mimeType='text/csv').execute().decode('utf-8')
        
        elif data_mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # NEW: Word .docx support
            print("   Format: Word .docx → Extracting text with python-docx")
            data_content = _read_docx_from_drive(data_file_id, drive_service)
        
        elif data_mime_type == 'application/msword':
            # Legacy .doc format
            print("   Format: Legacy Word .doc → Converting via Google Drive")
            # Google Drive can convert old .doc to text
            try:
                data_content = drive_service.files().export(fileId=data_file_id, mimeType='text/plain').execute().decode('utf-8')
            except:
                return f"Error: Cannot read legacy .doc format. Please convert to .docx or upload as Google Doc."
        
        elif data_mime_type == 'application/pdf':
            # PDF support (requires additional library)
            print("   Format: PDF → Not directly supported")
            return f"Error: PDF data files are not supported. Please convert PDF to text or Word format first."
        
        else:
            return f"Error: Unsupported data file type: {data_mime_type}\n\nSupported formats:\n- Google Docs\n- Word (.docx)\n- Text files (.txt)\n- CSV files\n- Google Sheets"
        
        print(f" Read {len(data_content)} characters from data file")
        
        # Step 2: Analyze template for placeholders
        extractor = DocumentFormatExtractor(credentials_dict)
        structure = extractor.extract_document_structure(template_file_id)
        
        if "error" in structure:
            return f"Error analyzing template: {structure['error']}"
        
        placeholders = extractor.identify_placeholders(structure)
        print(f" Found placeholders in template: {placeholders}")
        
        # Step 3: Parse data content into placeholder values
        placeholder_values = {}
        
        if not placeholders:
            placeholder_values = {"CONTENT": data_content}
            print("ℹ No placeholders found in template, using data as CONTENT")
        else:
            # Parse data as key-value pairs
            lines = data_content.strip().split('\n')
            
            for line in lines:
                separator = None
                if ':' in line:
                    separator = ':'
                elif '=' in line:
                    separator = '='
                elif '|' in line:
                    separator = '|'
                
                if separator:
                    parts = line.split(separator, 1)
                    if len(parts) == 2:
                        key = parts[0].strip().upper().replace(' ', '_').replace('-', '_')
                        value = parts[1].strip()
                        
                        # Match to placeholders
                        for placeholder in placeholders:
                            if key == placeholder or key in placeholder or placeholder in key:
                                placeholder_values[placeholder] = value
                                print(f" Matched: [{placeholder}] = {value[:50]}{'...' if len(value) > 50 else ''}")
                                break
            
            # Fill missing placeholders
            for placeholder in placeholders:
                if placeholder not in placeholder_values:
                    if placeholder in ['CONTENT', 'DATA', 'TEXT']:
                        placeholder_values[placeholder] = data_content
                        print(f" ℹ Using full data for [{placeholder}]")
                    elif placeholder == 'DATE':
                        from datetime import datetime
                        placeholder_values[placeholder] = datetime.now().strftime("%B %d, %Y")
                        print(f" ℹ Auto-filled [{placeholder}] with current date")
                    else:
                        placeholder_values[placeholder] = ""
                        print(f" No data for [{placeholder}], leaving empty")
        
        print(f" Final placeholder mapping: {list(placeholder_values.keys())}")
        
        # Step 4: Create document from template
        print(f" Creating document '{new_title}' from template...")
        result = extractor.create_from_template(
            template_document_id=template_file_id,
            new_title=new_title,
            placeholder_values=placeholder_values
        )
        
        if not result.get("success"):
            return f"Error creating document: {result.get('error')}"
        
        print(f" Document created: {result['document_id']}")
        
        # Step 5: Replace blank lines
        print(f" Checking for blank line placeholders...")
        _replace_blank_lines_in_document(
            document_id=result['document_id'],
            placeholder_values=placeholder_values,
            credentials_dict=credentials_dict
        )
        
        # Step 6: Export as PDF if requested
        if output_format == "pdf":
            print(" Exporting as PDF...")
            try:
                pdf_content = drive_service.files().export(
                    fileId=result['document_id'],
                    mimeType='application/pdf'
                ).execute()
                
                doc_metadata = drive_service.files().get(
                    fileId=result['document_id'],
                    fields='parents'
                ).execute()
                
                parents = doc_metadata.get('parents', [])
                
                pdf_metadata = {
                    'name': f"{new_title}.pdf",
                    'mimeType': 'application/pdf',
                    'parents': parents
                }
                
                from googleapiclient.http import MediaInMemoryUpload
                pdf_file = drive_service.files().create(
                    body=pdf_metadata,
                    media_body=MediaInMemoryUpload(pdf_content, mimetype='application/pdf'),
                    fields='id, webViewLink'
                ).execute()
                
                response = f"Document created and exported as PDF.\n\n"
                response += f"Title: {new_title}.pdf\n"
                response += f"PDF ID: {pdf_file['id']}\n"
                response += f"PDF URL: {pdf_file['webViewLink']}\n"
                response += f"Google Doc ID: {result['document_id']}\n"
                response += f"Google Doc URL: {result['url']}\n"
                
                return response
                
            except Exception as pdf_error:
                print(f"PDF export failed: {pdf_error}")
        
        response = f"Document created successfully.\n\n"
        response += f"Title: {result['title']}\n"
        response += f"Document ID: {result['document_id']}\n"
        response += f"URL: {result['url']}\n"
        
        return response
    
    except Exception as error:
        import traceback
        traceback.print_exc()
        return f"Error creating document: {error}"

# Add this function at the TOP of your tools.py file (after imports)
# This should go right after your imports and before get_google_service()

def _read_docx_from_drive(file_id: str, drive_service) -> str:
    """
    Read text content from a .docx file stored in Google Drive
    
    Args:
        file_id: Google Drive file ID
        drive_service: Authenticated Google Drive service
    
    Returns:
        Extracted text content from the Word document
    """
    try:
        import io
        from docx import Document
        
        # Download the file content
        request = drive_service.files().get_media(fileId=file_id)
        file_content = io.BytesIO()
        
        from googleapiclient.http import MediaIoBaseDownload
        downloader = MediaIoBaseDownload(file_content, request)
        
        done = False
        while not done:
            status, done = downloader.next_chunk()
        
        # Reset file pointer to beginning
        file_content.seek(0)
        
        # Parse with python-docx
        doc = Document(file_content)
        
        # Extract all text from paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                text_content.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        full_text = '\n'.join(text_content)
        print(f" Extracted {len(full_text)} characters from .docx file")
        
        return full_text
        
    except ImportError:
        raise ImportError(
            "python-docx library not installed. "
            "Install with: pip install python-docx"
        )
    except Exception as e:
        raise Exception(f"Error reading .docx file: {str(e)}")


def _read_file_content(file_path: str) -> str:
    """
    Read text content from a local file. Supports PDF (via pdfplumber) and
    plain-text formats (.txt, .html, .md, .json, .xml, .csv).

    No character cap — returns full content for direct Google Docs API insertion.
    Raises ValueError for unsupported file types.
    """
    import os

    if not file_path or not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber is required for PDF reading. Install with: pip install pdfplumber")

        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
        content = "\n".join(text_parts)
        if not content.strip():
            raise ValueError(f"PDF file is empty or could not be parsed: {file_path}")
        return content

    if ext in (".txt", ".html", ".md", ".json", ".xml", ".csv", ".log"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        if not content.strip():
            raise ValueError(f"File is empty: {file_path}")
        return content

    raise ValueError(
        f"Unsupported file type '{ext}' for direct reading. "
        f"Supported: .pdf, .txt, .html, .md, .json, .xml, .csv"
    )


def _create_doc_with_content_impl(
    title: str,
    credentials_dict: Dict,
    text: str = None,
    file_path: str = None,
    folder_id: str = None,
) -> Dict[str, Any]:
    """
    Create a new Google Doc and populate it with content in a single operation.

    Accepts either `text` (inline string) or `file_path` (local file to read).
    If both are provided, `file_path` takes precedence.

    folder_id: Optional Drive folder ID to place the new doc in. Passed through
    to `_create_google_doc_impl` which reparents via the Drive API after creation.

    Returns a structured dict (not a string) for direct-dispatch consumption.
    """
    import re

    if not text and not file_path:
        return {
            "success": False,
            "error": "Either 'text' or 'file_path' must be provided",
        }

    content = text or ""
    if file_path:
        try:
            content = _read_file_content(file_path)
            print(f"Read {len(content)} characters from {file_path}")
        except Exception as e:
            return {"success": False, "error": f"Failed to read file: {e}"}

    create_result = _create_google_doc_impl(title, credentials_dict, folder_id=folder_id)

    if "error" in create_result.lower() and "successfully" not in create_result.lower():
        return {"success": False, "error": create_result}

    doc_id_match = re.search(r"ID: ([a-zA-Z0-9_-]+)", create_result)
    if not doc_id_match:
        return {"success": False, "error": f"Could not parse document ID from: {create_result}"}

    doc_id = doc_id_match.group(1)
    doc_url = f"https://docs.google.com/document/d/{doc_id}/edit"

    folder_moved = "Folder moved: yes" in create_result
    folder_move_error = None
    if folder_id and not folder_moved:
        err_match = re.search(r"Folder move error: (.+)", create_result)
        if err_match:
            folder_move_error = err_match.group(1).strip()

    add_result = _add_text_to_doc_impl(doc_id, content, credentials_dict)

    if "error" in add_result.lower() and "successfully" not in add_result.lower():
        return {
            "success": False,
            "error": f"Document created (ID: {doc_id}) but failed to add content: {add_result}",
            "document_id": doc_id,
            "document_url": doc_url,
            "folder_id": folder_id,
            "folder_moved": folder_moved,
        }

    return {
        "success": True,
        "document_id": doc_id,
        "document_url": doc_url,
        "title": title,
        "text_length": len(content),
        "folder_id": folder_id,
        "folder_moved": folder_moved,
        "folder_move_error": folder_move_error,
    }


def _add_text_from_file_impl(
    document_id: str,
    file_path: str,
    credentials_dict: Dict,
) -> Dict[str, Any]:
    """
    Read a local file and add its text content to an existing Google Doc.

    Returns a structured dict (not a string) for direct-dispatch consumption.
    """
    document_id = _extract_document_id(document_id)
    try:
        content = _read_file_content(file_path)
        print(f"Read {len(content)} characters from {file_path}")
    except Exception as e:
        return {"success": False, "error": f"Failed to read file: {e}"}

    add_result = _add_text_to_doc_impl(document_id, content, credentials_dict)

    if "error" in add_result.lower() and "successfully" not in add_result.lower():
        return {"success": False, "error": add_result}

    doc_url = f"https://docs.google.com/document/d/{document_id}/edit"
    return {
        "success": True,
        "document_id": document_id,
        "document_url": doc_url,
        "text_length": len(content),
    }